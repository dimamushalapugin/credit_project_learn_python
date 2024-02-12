import openpyxl
import docx
import os
import datetime
import requests

from bs4 import BeautifulSoup
from docx import Document
from dadata import Dadata
from docx.shared import Pt
from datetime import datetime as dt
from flask_login import current_user
from webapp.risk.logger import logging
from webapp.managers.parser_for_dkp import read_xlsx, number_to_words
from webapp.managers.merge_two_xlsx import merge_files
from webapp.config import DADATA_TOKEN, DADATA_BASE, dl_car_ip_path, dl_obor_ip_path, dl_car_ooo_path, dl_obor_ooo_path
from webapp.risk.mongo_db import MongoDB
from openpyxl.worksheet.datavalidation import DataValidation
from pathlib import Path


def start_filling_application(inn_leasee, path_application, inn_seller1, inn_seller2, inn_seller3, inn_seller4):
    mongo = MongoDB(current_user)
    logging.info(f"({current_user}) Этап 1.")
    temporary_path = Path("webapp") / "static" / "temporary"
    path_for_download = Path("static") / "temporary"
    ip_or_kfh = 'Нет'
    type_business = ''
    full_krakt_name_leasee = ''
    main_activity_leasee = ''
    ogrn_leasee = ''
    okpo_leasee = ''
    okato_leasee = ''
    date_regist = ''
    ustav_capital = ''
    inn_kpp_leasee = ''
    address_leasee = ''
    formatted_name_leader_leasee = ''
    fio_list = ''
    inn_list = ''
    dolya_list = ''
    full_name_leasee = ''
    leader_leasee = ''
    fio_leader = ''
    phone_leasee = ''
    email_leasee = ''
    krakt_name_seller1 = ''
    krakt_name_seller2 = ''
    krakt_name_seller3 = ''
    krakt_name_seller4 = ''
    address_seller1 = ''
    address_seller2 = ''
    address_seller3 = ''
    address_seller4 = ''
    inn_dir_leasee = ''

    def parser_info_leasee(inn_leasee):
        logging.info(f"({current_user}) Этап 2.")
        nonlocal ip_or_kfh, type_business, inn_dir_leasee, krakt_name_seller1, address_seller1, krakt_name_seller2, address_seller2, krakt_name_seller3, address_seller3, krakt_name_seller4, address_seller4, full_krakt_name_leasee, main_activity_leasee, ogrn_leasee, okpo_leasee, okato_leasee, date_regist, ustav_capital, inn_kpp_leasee, address_leasee, formatted_name_leader_leasee, fio_list, inn_list, dolya_list, full_name_leasee, leader_leasee, fio_leader, phone_leasee, email_leasee

        logging.info(f"({inn_leasee})")

        dadata = Dadata(DADATA_TOKEN)
        result = dadata.find_by_id("party", inn_leasee)
        # logging.info(f"{result}")

        ip_or_kfh = 'Нет'
        if result[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
            ip_or_kfh = 'Да'

        full_name_leasee = result[0]['data']['name']['full_with_opf']

        full_krakt_name_leasee = result[0]['data']['name']['short_with_opf']

        if ip_or_kfh == 'Нет':
            inn_kpp_leasee = inn_leasee + '/' + result[0]['data']['kpp']
            leader_leasee = result[0]['data']['management']['post']
            fio_leader = result[0]['data']['management']['name']
        else:
            inn_kpp_leasee = inn_leasee
            type_business = result[0]['data']['opf']['short']
            leader_leasee = result[0]['data']['opf']['short']
            if result[0]['data']['opf']['short'] == 'ИП':
                leader_leasee = 'Индивидуальный предприниматель'
            else:
                leader_leasee = 'Глава'
            fio_leader = result[0]['data']['name']['full']

        last_name, first_name, patronymic_name = fio_leader.split()

        # Get the initial of the first name
        first_name_initial = first_name[0]

        # Get the initial of the patronymic name
        patronymic_initial = patronymic_name[0]

        # Combine the last name and initials in the desired format
        formatted_name_leader_leasee = f'{last_name} {first_name_initial}.{patronymic_initial}.'

        address_leasee = result[0]['data']['address']['unrestricted_value']

        ogrn_leasee = result[0]['data']['ogrn']

        okato_leasee = result[0]['data']['okato']

        okpo_leasee = result[0]['data']['okpo']

        # Значение registration_date в миллисекундах
        registration_date_ms = result[0]['data']['ogrn_date']
        # Преобразуем значение в объект datetime
        registration_date = datetime.datetime.fromtimestamp(registration_date_ms / 1000)
        # Преобразование строки в объект datetime
        date_time_obj = datetime.datetime.strptime(str(registration_date), "%Y-%m-%d %H:%M:%S")
        # Форматирование даты в нужный формат
        date_regist = date_time_obj.strftime("%d.%m.%Y")

        def parse_client_bs4(ogrn):
            nonlocal fio_list, inn_list, dolya_list, inn_dir_leasee, phone_leasee, email_leasee, main_activity_leasee, ustav_capital
            url = f'https://vbankcenter.ru/contragent/{ogrn}'
            url2 = f'https://vbankcenter.ru/contragent/search?searchStr={ogrn}'

            response = requests.get(url)
            response2 = requests.get(url2)

            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                soup2 = BeautifulSoup(response2.text, 'html.parser')

                if len(str(ogrn)) == 13:

                    try:
                        fio_list = list(map(lambda x: x.text, soup.find_all('h5', class_='text-base font-bold')))
                    except Exception as ex:
                        fio_list = ''
                        logging.info(ex, exc_info=True)

                    dolya_list = []
                    try:
                        for elem in soup.find_all('p', class_='text-base m-0 text-premium-600', string='Доля:'):
                            dolya_list.append(
                                "".join(elem.find_next('p', class_='text-base m-0 ml-1.5').text.split()[-1]).replace(
                                    '%',
                                    '').replace(
                                    '(', '').replace(')', ''))
                    except Exception as ex:
                        logging.info(ex, exc_info=True)

                    inn_list = []
                    try:
                        for elem in soup.find_all('p', class_='text-base m-0 text-premium-600', string='ИНН:'):
                            inn_list.append(elem.find_next('a', class_='flex text-base m-0 ml-1.5').text)
                    except Exception as ex:
                        logging.info(ex, exc_info=True)

                    try:
                        inn_dir_leasee = soup.find('p', class_='mb-1 whitespace-nowrap pr-6 text-premium-600').find('a',
                                                                                                                    class_='text-blue').text
                    except Exception as ex:
                        inn_dir_leasee = ''
                        logging.info(ex, exc_info=True)

                    try:
                        for elem in soup.find(class_='requisites-info-badge font-bold mb-1').find_next('div',
                                                                                                       class_='flex items-baseline mt-1').find(
                            class_='gweb-copy relative inline-block mb-0 py-0 copy-available z-10 cursor-pointer copy-right-padding'):
                            phone_leasee = elem.get_text(strip=True)
                    except Exception as ex:
                        phone_leasee = ''
                        logging.info('Нет телефона')
                        logging.info(ex, exc_info=True)

                    try:
                        for elem in soup.find(class_='requisites-info-badge font-bold mb-1').find_next('div',
                                                                                                       class_='flex items-baseline mt-1').find_next(
                            'div', class_='flex items-baseline mt-1').find('a'):
                            email_leasee = elem.get_text(strip=True)
                    except Exception as ex:
                        email_leasee = ''
                        logging.info('Нет email')
                        logging.info(ex, exc_info=True)

                    try:
                        main_activity_leasee = soup.find('a', class_='inline-block mt-1').get_text(strip=True)
                    except Exception as ex:
                        main_activity_leasee = ''
                        logging.info(ex, exc_info=True)

                    try:
                        number_element = soup2.find('span', class_='inline-block pr-2 text-premium-600 lg:mr-4 xl:mr-0',
                                                    string='Уставный капитал:').find_next_sibling('span')
                        number_text = number_element.text.strip()
                        ustav_capital = float(''.join(filter(lambda x: x.isdigit() or x == '.', number_text)))
                    except Exception as ex:
                        ustav_capital = ''
                        logging.info(ex, exc_info=True)

                else:
                    inn_dir_leasee = inn_leasee
                    phone_leasee = ''
                    email_leasee = ''
                    ustav_capital = ''
                    try:
                        main_activity_leasee = soup.find('a', class_='inline-block mt-1').get_text(strip=True)
                    except Exception as ex:
                        main_activity_leasee = ''
                        logging.info(ex, exc_info=True)

            else:
                logging.info(f"Ошибка при запросе: {response.status_code}")
                fio_list = ''
                dolya_list = []
                inn_list = []
                inn_dir_leasee = ''
                phone_leasee = ''
                email_leasee = ''
                ustav_capital = ''
                main_activity_leasee = ''

        parse_client_bs4(ogrn_leasee)

        logging.info(f"({current_user}) Этап 3.")

        result_seller1 = dadata.find_by_id("party", inn_seller1)
        # print(result_seller1)

        krakt_name_seller1 = result_seller1[0]['data']['name']['short_with_opf']
        # print(f'Здесь брать краткое наименование Продавца №1 {krakt_name_seller1}')

        address_seller1 = result_seller1[0]['data']['address']['unrestricted_value']
        # print(f'Здесь брать юр адрес Продавца №1 {address_seller1}')

        # print(len(inn_seller2))
        if len(inn_seller2) >= 1 and inn_seller2 is not None:
            result_seller2 = dadata.find_by_id("party", inn_seller2)
            # print(result_seller2)

            krakt_name_seller2 = result_seller2[0]['data']['name']['short_with_opf']
            # print(f'Здесь брать краткое наименование Продавца №2 {krakt_name_seller2}')

            address_seller2 = result_seller2[0]['data']['address']['unrestricted_value']
            # print(f'Здесь брать юр адрес Продавца №2 {address_seller2}')

        if len(inn_seller3) >= 1 and inn_seller3 is not None:
            result_seller3 = dadata.find_by_id("party", inn_seller3)
            # print(result_seller3)

            krakt_name_seller3 = result_seller3[0]['data']['name']['short_with_opf']
            # print(f'Здесь брать краткое наименование Продавца №3 {krakt_name_seller3}')

            address_seller3 = result_seller3[0]['data']['address']['unrestricted_value']
            # print(f'Здесь брать юр адрес Продавца №3 {address_seller3}')

        if len(inn_seller4) >= 1 and inn_seller4 is not None:
            result_seller4 = dadata.find_by_id("party", inn_seller4)
            # print(result_seller4)

            krakt_name_seller4 = result_seller4[0]['data']['name']['short_with_opf']
            # print(f'Здесь брать краткое наименование Продавца №4 {krakt_name_seller4}')

            address_seller4 = result_seller4[0]['data']['address']['unrestricted_value']
            # print(f'Здесь брать юр адрес Продавца №4 {address_seller4}')

    def zapolnenie_zayvki_ankety(inn_leasee, path_application, inn_dir_leasee):
        logging.info(f"({current_user}) Этап 4.")
        try:
            # сейчас будем заполнять заявку, вносить данные по лп
            wb = openpyxl.load_workbook(rf'{path_application}')
            # заполняем страницу Заявление
            sheet_zayavlenie = wb['Заявление']
            # Создаем список валидации данных
            dv = DataValidation(type="list", formula1='"Руб с НДС, Китайский юань с НДС, Доллар с НДС"', showDropDown=True)
            # Apply the DataValidation object to the cell
            sheet_zayavlenie.add_data_validation(dv)
            dv.add(sheet_zayavlenie['Q21'])

            sheet_zayavlenie['A5'].value = full_name_leasee
            sheet_zayavlenie['D6'].value = inn_kpp_leasee
            sheet_zayavlenie['B1'].value = dt.today().strftime(f"%d.%m.%Y")
            if ip_or_kfh == 'Да':
                sheet_zayavlenie['D6'].value = inn_leasee
            # print(sheet_zayavlenie['A6'].value)
            # print(sheet_zayavlenie['C7'].value)
            sheet_zayavlenie['H6'].value = address_leasee
            # print(sheet_zayavlenie['E7'].value)

            sheet_zayavlenie['A10'].value = krakt_name_seller1
            sheet_zayavlenie['C11'].value = inn_seller1
            sheet_zayavlenie['G11'].value = address_seller1
            if len(inn_seller2) >= 1 and inn_seller2 is not None:
                sheet_zayavlenie['A12'].value = krakt_name_seller2
                sheet_zayavlenie['C13'].value = inn_seller2
                sheet_zayavlenie['G13'].value = address_seller2
            if len(inn_seller3) >= 1 and inn_seller3 is not None:
                sheet_zayavlenie['A14'].value = krakt_name_seller3
                sheet_zayavlenie['C15'].value = inn_seller3
                sheet_zayavlenie['G15'].value = address_seller3
            if len(inn_seller4) >= 1 and inn_seller4 is not None:
                sheet_zayavlenie['A16'].value = krakt_name_seller4
                sheet_zayavlenie['C17'].value = inn_seller4
                sheet_zayavlenie['G17'].value = address_seller4

            counter_1 = 24
            for number in range(25, sheet_zayavlenie.max_row + 2):
                counter_1 += 1
                if sheet_zayavlenie[
                    f'B{number}'].value == 'Место эксплуатации предмета лизинга (для автотранспорта место стоянки/хранения) полный фактический адрес:':
                    sheet_zayavlenie[f'A{number + 1}'].value = address_leasee
                    # print(sheet_zayavlenie[f'A{number + 1}'].value)
                if sheet_zayavlenie[f'A{number}'].value == '(должность руководителя организации Заявителя)':
                    sheet_zayavlenie[f'A{number - 1}'].value = leader_leasee
                    # print(sheet_zayavlenie[f'B{number - 1}'].value)
                if sheet_zayavlenie[f'O{number}'].value == '(расшифровка подписи)':
                    sheet_zayavlenie[f'O{number - 1}'].value = formatted_name_leader_leasee
                    # print(sheet_zayavlenie[f'H{number - 1}'].value)

                # заполнение поручителей, автоматом поставил всех учредов
                if ip_or_kfh != 'Да':
                    if sheet_zayavlenie[f'B{number}'].value == 'Наименование поручителя ⃰ ':
                        for row_num, fio in enumerate(fio_list, start=1):
                            sheet_zayavlenie.cell(row=row_num + counter_1, column=2, value=fio)
                        # print(sheet_zayavlenie.cell(row=row_num + counter_1, column=2, value=fio).value)
                    if sheet_zayavlenie[f'J{number}'].value == 'ИНН':
                        for row_num, inn in enumerate(inn_list, start=1):
                            sheet_zayavlenie.cell(row=row_num + counter_1, column=10, value=inn)

                        # print(sheet_zayavlenie.cell(row=row_num + counter_1, column=6, value=inn).value)
                if ip_or_kfh == 'Да':
                    if sheet_zayavlenie[f'B{number}'].value == 'Наименование поручителя ⃰ ':
                        sheet_zayavlenie[f'B{number + 1}'].value = fio_leader
                        # print(sheet_zayavlenie[f'B{number + 1}'].value)
                    if sheet_zayavlenie[f'J{number}'].value == 'ИНН':
                        sheet_zayavlenie[f'J{number + 1}'].value = inn_leasee
                        # print(sheet_zayavlenie[f'F{number + 1}'].value)
                    if sheet_zayavlenie[f'A{number}'].value == '(должность руководителя организации Заявителя)':
                        if type_business == 'Индивидуальный предприниматель' or type_business == 'ИП':
                            sheet_zayavlenie[f'A{number - 1}'].value = type_business
                        else:
                            sheet_zayavlenie[f'A{number - 1}'].value = 'Глава'

            # заполняем страницу Анкета Стр.1
            sheet_anketa_1_list = wb['Анкета_Стр.1']
            sheet_anketa_1_list['F7'].value = ogrn_leasee
            # print(sheet_anketa_1_list['F7'].value)
            sheet_anketa_1_list['H7'].value = okato_leasee
            # print(sheet_anketa_1_list['H7'].value)
            sheet_anketa_1_list['J7'].value = okpo_leasee
            # print(sheet_anketa_1_list['J7'].value)
            sheet_anketa_1_list['E8'].value = date_regist
            # print(sheet_anketa_1_list['E8'].value)
            sheet_anketa_1_list['J9'].value = ustav_capital
            # print(sheet_anketa_1_list['J9'].value)
            sheet_anketa_1_list['A6'].value = full_krakt_name_leasee
            # print(sheet_anketa_1_list['A6'].value)
            bank_details = mongo.read_mongodb_bank_details(inn_leasee)
            director_details = mongo.read_mongodb_director_details(inn_dir_leasee)
            logging.info(bank_details)
            logging.info(director_details)
            if bank_details:
                sheet_anketa_1_list['G40'].value = bank_details.get('bank')
                sheet_anketa_1_list['B41'].value = bank_details.get('check_account')
                sheet_anketa_1_list['F41'].value = bank_details.get('cor_account')
                sheet_anketa_1_list['I41'].value = bank_details.get('bik')
            if director_details:
                sheet_anketa_1_list['D24'].value = director_details.get('date_of_birth')
                sheet_anketa_1_list['F24'].value = director_details.get('place_of_birth')
                sheet_anketa_1_list['D28'].value = director_details.get('passport_series')
                sheet_anketa_1_list['D29'].value = director_details.get('passport_id')
                sheet_anketa_1_list['F28'].value = director_details.get('issued_by')
                sheet_anketa_1_list['F29'].value = director_details.get('issued_when')
                sheet_anketa_1_list['D30'].value = director_details.get('department_code')
                sheet_anketa_1_list['D31'].value = director_details.get('address_reg')
                sheet_anketa_1_list['E32'].value = director_details.get('address_fact')

            counter_2_anketa = 7
            for number in range(8, sheet_anketa_1_list.max_row + 2):
                counter_2_anketa += 1
                # заполняем инфу по учредам, не более 4-х должно быть
                if ip_or_kfh == 'Нет':
                    if sheet_anketa_1_list[
                        f'B{number}'].value == 'полное наименование акционера/участника/члена/товарища с указанием организационно-правовой формы (полностью Ф.И.О. для физических лиц):':
                        for row_num, fio in enumerate(fio_list, start=1):
                            sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=2, value=fio)
                        # print(sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=2, value=fio).value)
                    if sheet_anketa_1_list[f'G{number}'].value == 'ИНН':
                        # print(counter_2_anketa)
                        for row_num, inn in enumerate(inn_list, start=1):
                            sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=7, value=inn)
                        # print(sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=7, value=inn).value)
                    if sheet_anketa_1_list[f'I{number}'].value == 'доля в уставном капитале, в %':
                        for row_num, dolya in enumerate(dolya_list, start=1):
                            # dolya_value = float(dolya)
                            sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=9, value=dolya)
                        # print(sheet_anketa_1_list.cell(row=row_num + counter_2_anketa + 1, column=9).value)
                    if sheet_anketa_1_list[f'G{number}'].value == 'ИНН:':
                        sheet_anketa_1_list[f'H{number}'].value = inn_dir_leasee
                if sheet_anketa_1_list[f'A{number}'].value == '1.7         Телефон:':
                    sheet_anketa_1_list[f'C{number}'].value = phone_leasee
                    # print(sheet_anketa_1_list[f'C{number}'].value)
                if sheet_anketa_1_list[f'E{number}'].value == '1.8 Эл. почта:':
                    sheet_anketa_1_list[f'F{number}'].value = email_leasee
                    # print(sheet_anketa_1_list[f'F{number}'].value)
                if sheet_anketa_1_list[f'B{number}'].value == 'ФИО:':
                    sheet_anketa_1_list[f'C{number}'].value = fio_leader
                    # print(sheet_anketa_1_list[f'C{number}'].value)
                if sheet_anketa_1_list[f'B{number}'].value == 'ОКВЭД с расшифровкой:':
                    sheet_anketa_1_list[f'E{number}'].value = main_activity_leasee
                    # print(sheet_anketa_1_list[f'E{number}'].value)
                if ip_or_kfh == 'Да':
                    if sheet_anketa_1_list[f'G{number}'].value == 'ИНН:':
                        sheet_anketa_1_list[f'H{number}'].value = inn_leasee

            if bank_details:
                sheet_anketa_1_list['C21'].value = bank_details.get('phone')
                sheet_anketa_1_list['F21'].value = bank_details.get('email')

            application_filename = temporary_path / fr'Заявка с заключением {inn_leasee} (read).xlsx'
            wb.save(application_filename)
            logging.info(f"({current_user}) Запускаем объединение файлов")
            application_filename_download = path_for_download / fr'Заявка с заключением {inn_leasee}.xlsx'
            xlsx_name_read = fr'Заявка с заключением {inn_leasee} (read).xlsx'
            xlsx_name = fr'Заявка с заключением {inn_leasee}.xlsx'
            merge_files(xlsx_name_read, xlsx_name)
            logging.info(f"({current_user}) Все успешно сохранилось!")

            return application_filename_download
        except Exception as ex:
            logging.info(ex, exc_info=True)
            raise ValueError

    parser_info_leasee(inn_leasee)  # берет инфу из инета по лизингополучателю
    application_filename = zapolnenie_zayvki_ankety(inn_leasee,
                                                    path_application,
                                                    inn_dir_leasee)  # заполняет эксель данными из инета по лизингополучателю

    return application_filename


def start_filling_agreement(inn_leasee, path_application, path_graphic, signatory, investor, currency_list,
                            who_is_insure, grafic, pl, number_dl, inn_seller, type_pl):
    dadata = Dadata(DADATA_TOKEN)
    result = dadata.find_by_id("party", inn_leasee)

    full_krakt_name_leasee = result[0]['data']['name']['short_with_opf'].replace('"', '')

    dir_path = Path('webapp') / 'static' / 'agreements' / f'{full_krakt_name_leasee} {inn_leasee}' / f'{dt.today().strftime(f"%d.%m.%Y")}'
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

    data_xlsx = read_xlsx(path_application, pl)

    formatted_name_leader_leasee = data_xlsx[17]
    full_name_leasee = data_xlsx[22]
    full_krakt_name_leasee = data_xlsx[8]
    inn_kpp_leasee = data_xlsx[21]
    address_leasee = data_xlsx[20]
    address_leasee_expluatazia = data_xlsx[19]
    predmet_lizinga = pl
    inn_seller_list = inn_seller
    price_predmet_lizinga = data_xlsx[15]
    ogrn_leasee = data_xlsx[13]
    okato_leasee = data_xlsx[12]
    okpo_leasee = data_xlsx[11]
    date_regist = data_xlsx[10]
    ustav_capital = data_xlsx[9]
    phone_leasee = data_xlsx[7]
    email_leasee = data_xlsx[6]
    fio_leader = data_xlsx[5]
    main_activity_leasee = data_xlsx[4]
    rekvizit_leasee_bank = data_xlsx[3]
    rekvizit_leasee_shet = data_xlsx[2]
    rekvizit_leasee_cs_shet = data_xlsx[1]
    rekvizit_leasee_bik = data_xlsx[0]
    leader_leasee = data_xlsx[18]

    # Чтение заявки и анкеты

    vikup = ''
    vigodo = ''

    input_raschet_path = rf'{path_graphic}'

    def create_dl_dkp(inn_leasee, path_application, path_graphic, signatory, investor, currency_list, who_is_insure,
                      grafic, pl, number_dl, inn_seller):
        nonlocal leader_leasee, vikup, vigodo
        if signatory == 'Каюмов А. Д.':
            a_lkmb = 'Директор'
            lkmb_podpisant = 'Каюмов А. Д.'
            preambula_dolj_lkmb = 'Директора'
            preambula_fio_lkmb = 'Каюмова Айрата Дамировича'
            doverka_ustav_list = 'Устава'
            deystvuysh_list = 'действующего'
        elif signatory == 'Габдрахманов Р. Р.':
            a_lkmb = 'Заместитель директора'
            lkmb_podpisant = 'Габдрахманов Р. Р.'
            preambula_dolj_lkmb = 'Заместителя директора'
            preambula_fio_lkmb = 'Габдрахманова Рината Рафаэлевича'
            doverka_ustav_list = 'Доверенности от «17» марта 2020 года, удостоверенной Мальченковой  Евгенией Николаевной, нотариусом Казанского нотариального округа Республики Татарстан, зарегистрированной в реестре нотариальных действий за № 16/64-н/16-2020-7-317(бланк 16 АА 5665323)'
            deystvuysh_list = 'действующего'
        else:
            a_lkmb = 'Заместитель директора по финансам'
            lkmb_podpisant = 'Хасанова Д. Р.'
            preambula_dolj_lkmb = 'Заместителя директора по финансам'
            preambula_fio_lkmb = 'Хасановой Динары Ринатовны'
            doverka_ustav_list = 'Доверенности от «17» марта 2020 года, удостоверенной Мальченковой  Евгенией Николаевной, нотариусом Казанского нотариального округа Республики Татарстан, зарегистрированной в реестре нотариальных действий за № 16/64-н/16-2020-7-317(бланк 16 АА 5665323)'
            deystvuysh_list = 'действующей'
        if leader_leasee.upper() == 'директор'.upper():
            leader_leasee_rod_padezh = 'Директора'
        elif leader_leasee.upper() == 'генеральный директор'.upper():
            leader_leasee_rod_padezh = 'Генерального директора'
        elif leader_leasee.upper() == 'исполняющий обязанности директора'.upper():
            leader_leasee_rod_padezh = 'ИО директора'
        else:
            leader_leasee_rod_padezh = ''

        # запускается функция по замене ФИО подписанта лизингополучателя
        put_padezh_podpisant = ''

        def rod_padezh_fio_leader(fio):
            put_padezh_podpisant = DADATA_BASE.clean("name", fio)
            return put_padezh_podpisant

        rod_padezh_fio_leader = rod_padezh_fio_leader(data_xlsx[5])
        try:
            put_padezh_podpisant_rg = rod_padezh_fio_leader['result_genitive']
        except:
            put_padezh_podpisant_rg = ''
        # print(f'123 {put_padezh_podpisant_rg}')
        doverka_ustav_leasee = 'Устава'
        for elem in full_name_leasee.split():
            if elem in ['Индивидуальный', 'предприниматель', 'хозяйства']:
                doverka_ustav_leasee = f'Свидетельства о государственной регистрации физического лица в качестве индивидуального предпринимателя от {date_regist}, ОГРНИП {ogrn_leasee}'

        deystvuysh_list_leasee = 'действующей'
        try:
            if rod_padezh_fio_leader['gender'] == 'М':
                deystvuysh_list_leasee = 'действующего'
                if result[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_leasee = 'действующий'
                imenyemoe = 'именуемый'
            else:
                deystvuysh_list_leasee = 'действующая'
                imenyemoe = 'именуемая'
        except:
            try:
                if result[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_leasee = 'действующая'
            except:
                deystvuysh_list_leasee = 'действующей'
            imenyemoe = 'именуемая'
        # deystvuysh_list_leasee = 'действующей' if fio_leader.split()[0][-1].lower() == 'а' else 'действующего'

        if investor in ['ПАО «МКБ»', 'ООО «ЛКМБ-РТ»']:
            vigodo = 'Лизингодатель'
        else:
            vigodo = investor

        # print('1912132')
        r_chet_lkmb = ''
        bank_rekv_lkmb = ''
        kor_chet_lkmb = ''
        bik_lkmb = ''

        if investor.upper() == 'ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»'.upper():
            r_chet_lkmb = '40701810000990000052'
            bank_rekv_lkmb = 'ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»'
            kor_chet_lkmb = '30101810300000000176'
            bik_lkmb = '044525176'
        elif investor.upper() == 'ПАО «МКБ»'.upper():
            r_chet_lkmb = '40701810900760000034'
            bank_rekv_lkmb = 'ПАО «МОСКОВСКИЙ КРЕДИТНЫЙ БАНК»'
            kor_chet_lkmb = '30101810745250000659'
            bik_lkmb = '044525659'
        elif investor.upper() == 'ИНВЕСТТОРГБАНК АО'.upper():
            r_chet_lkmb = '40701810071010300002'
            bank_rekv_lkmb = 'ИНВЕСТТОРГБАНК АО'
            kor_chet_lkmb = '30101810645250000267'
            bik_lkmb = '044525267'
        elif investor.upper() == 'АО «АЛЬФА-БАНК»'.upper():
            r_chet_lkmb = '40701810129930000005'
            bank_rekv_lkmb = 'ФИЛИАЛ «НИЖЕГОРОДСКИЙ» АО «АЛЬФА-БАНК»'
            kor_chet_lkmb = '30101810200000000824'
            bik_lkmb = '042202824'
        elif investor.upper() == 'АО «ПЕРВОУРАЛЬСКБАНК»'.upper():
            r_chet_lkmb = '40701810000010055037'
            bank_rekv_lkmb = 'АО «ПЕРВОУРАЛЬСКБАНК»'
            kor_chet_lkmb = '30101810565770000402'
            bik_lkmb = '046577402'
        elif investor.upper() == 'АО «СОЛИД БАНК»'.upper():
            r_chet_lkmb = '40701810105040011167'
            bank_rekv_lkmb = 'Московский филиал АО «СОЛИД БАНК»'
            kor_chet_lkmb = '30101810845250000795'
            bik_lkmb = '044525795'
        elif investor.upper() == 'АО КБ «УРАЛ ФД»'.upper():
            r_chet_lkmb = '40701810000000000962'
            bank_rekv_lkmb = 'АО КБ «УРАЛ ФД»'
            kor_chet_lkmb = '30101810800000000790'
            bik_lkmb = '045773790'
        else:
            r_chet_lkmb = '40702810100020002464'
            bank_rekv_lkmb = 'ПАО «АК БАРС» БАНК г. Казань'
            kor_chet_lkmb = '30101810000000000805'
            bik_lkmb = '049205805'

        currency = ''
        if currency_list == 'Рубль':
            currency_test = 'рублей'
        elif currency_list == 'Китайский юань':
            currency_test = 'юаней'
        elif currency_list == 'Доллар США':
            currency_test = 'долларов США'

        wr_rub_usd = {'рублей': 'рублях', 'доллары США': 'долларах США', 'ЕВРО': 'ЕВРО', 'юань': 'юанях',
                      'китайский юань': 'китайских юанях', 'рубль': 'рублях', 'доллары': 'долларах',
                      'руб.': 'рублях',
                      'юани': 'юанях', 'юаней': 'юанях', 'долларов США': 'долларах США'}
        # print(currency)
        # print(wr_rub_usd)
        # print('1910')
        leader_leasee_pod = leader_leasee
        inn_kpp1 = 'ИНН/КПП'
        ogrnip = ''
        if result[0]['data']['opf']['short'] in ['ИП', 'ГКФХ', 'КФХ']:
            inn_kpp1 = 'ИНН'
            ogrnip = f'ОГРНИП {ogrn_leasee}'
            leader_leasee_pod = ''
            put_padezh_podpisant_rg = ''
        else:
            imenyemoe = 'именуемое'
        # print(f'123213 {put_padezh_podpisant_rg}')
        vikup = '1000'
        pl_entry = pl

        price_entry = price_predmet_lizinga
        if vikup == '1000':
            punkt_4_6 = '4.6. Выкупная цена предмета лизинга составляет 1 000,00 (Одна тысяча) рублей, в том числе НДС.'
        else:
            punkt_4_6 = '4.6. Выкупная цена предмета лизинга, равная остаточной стоимости этого предмета лизинга, рассчитывается в соответствии с действующим законодательством Российской Федерации с учетом согласованной нормы амортизации. При этом выкупные платежи, уплаченные Лизингополучателем в составе лизинговых платежей в соответствии с Приложением № 1 к настоящему договору в качестве выкупных платежей, принимаются в зачет оплаты выкупной цены предмета лизинга. В случае расторжения настоящего договора по вине Лизингополучателя, в случае отказа от приемки предмета лизинга уплаченные Лизингополучателем выкупные платежи возвращению не подлежат, а удерживаются в качестве штрафа.'

        if investor == 'ИНВЕСТТОРГБАНК АО':
            punkt_7_8 = '\n'.join([
                '7.8. Лизингополучатель уведомлен:                                                                       ',
                '- что предмет лизинга обременен залогом АО «Инвестторгбанк» по кредитному договору, в рамках которого',
                'осуществляется финансирование лизинговой сделки, и что права требования на получение платежей по договору',
                'лизинга переданы в залог Банку;                                                                                          ',
                '- что Лизингополучатель при получении соответствующего уведомления от Банка (залогодержателя) обязан',
                'исполнять свое обязательство по договору лизинга Залогодержателю. При этом уведомление направляется любым из',
                'способов, определенных договором залога.'])
        else:
            punkt_7_8 = 'УДАЛИТЬ'

        suma_dann = ''

        # print('191919')

        def chislo_propis():
            nonlocal suma_dann
            suma_chislo = price_entry

            suma_dann = number_to_words(str(suma_chislo), currency_list)
            # print(f'01010 {suma_dann}')

        chislo_propis()

        summa_dog_leas = ''
        suma_dann_dl = ''

        def chislo_propis_dl(vybor_grafic_list):
            nonlocal summa_dog_leas, suma_dann_dl
            book = openpyxl.load_workbook(input_raschet_path, data_only=True)
            sheet = book[vybor_grafic_list]
            summa_dog_leas = sheet['F7'].value

            # print(type(f'Сумма дл{summa_dog_leas}'))

            suma_dann_dl = number_to_words(str(summa_dog_leas), currency_list)

        chislo_propis_dl(grafic)

        try:
            book = openpyxl.load_workbook(input_raschet_path, data_only=True)
            sheet = book[grafic]
            b93 = sheet['B93'].value.strftime('%d.%m.%Y') if sheet['B93'].value is not None else ''
            f7 = f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.', ',') if \
                f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                    -3] == ',' else f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.', ',') + '0'
        except Exception as ex:
            logging.info(ex, exc_info=True)
            b93 = ''
            f7 = ''

        if currency_test == 'рублей':
            currency_test = ''
        months = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня',
                  7: 'июля', 8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}

        old_words = ["{{ a_lkmb }}", "{{ lkmb_podpisant }}", "{{ preambula_dolj_lkmb }}",
                     "{{ preambula_fio_lkmb }}",
                     "{{ deystvuysh }}", "{{ doverka_ustav }}", "{{ full_name_leasee }}",
                     "{{ leader_leasee }}", "{{ fio_leader }}",
                     "{{ doverka_ustav_leasee }}", "{{ inn_seller }}", "{{ vigodo }}", "{{ phone_leasee }}",
                     "{{ email_leasee }}", "{{ full_krakt_name_leasee }}",
                     "{{ address_leasee_expluatazia }}", "{{ investor }}", "{{ r_chet_lkmb }}",
                     "{{ bank_rekv_lkmb }}",
                     "{{ kor_chet_lkmb }}", "{{ bik_lkmb }}",
                     "{{ inn_kpp_leasee }}", "{{ address_leasee }}", "{{ formatted_name_leader_leasee }}",
                     "{{ rekvizit_leasee_bank }}", "{{ rekvizit_leasee_shet }}",
                     "{{ rekvizit_leasee_cs_shet }}", "{{ rekvizit_leasee_bik }}", "{{ deystvuysh_list_leasee }}",
                     "{{ put_padezh_podpisant_rg }}", "{{ leader_leasee_rod_padezh }}", "{{ pl_entry }}",
                     "{{ price_entry }}",
                     "{{ number_dl }}", "{{ currency_test }}", "{{ suma_dann[0] }}", "{{ dt.today().day }}",
                     "{{ months[dt.today().month] }}", "{{ dt.today().year }}", "{{ punkt_4_6 }}",
                     "{{ summa_dog_leas }}", "{{ punkt_7_8 }}", "{{ inn_kpp1 }}", "{{ ogrnip }}",
                     "{{ leader_leasee_pod }}", "{{ imenyemoe }}", "{{ F7copy }}", "{{ B93copy }}"]
        # ,
        new_words = [str(a_lkmb), str(lkmb_podpisant), str(preambula_dolj_lkmb), str(preambula_fio_lkmb),
                     str(deystvuysh_list),
                     str(doverka_ustav_list),
                     str(full_name_leasee), str(leader_leasee),
                     str(fio_leader), str(doverka_ustav_leasee), str(inn_seller),
                     str(vigodo), str(phone_leasee), str(email_leasee), str(full_krakt_name_leasee),
                     str(address_leasee_expluatazia),
                     str(investor),
                     str(r_chet_lkmb), str(bank_rekv_lkmb), str(kor_chet_lkmb),
                     str(bik_lkmb), str(inn_kpp_leasee), str(address_leasee), str(formatted_name_leader_leasee),
                     str(rekvizit_leasee_bank),
                     str(rekvizit_leasee_shet), str(rekvizit_leasee_cs_shet),
                     str(rekvizit_leasee_bik), str(deystvuysh_list_leasee),
                     str(put_padezh_podpisant_rg),
                     str(leader_leasee_rod_padezh),
                     str(pl_entry), f'{price_entry:,.2f}'.replace(',', ' ').replace('.', ','), str(number_dl),
                     str(currency_test), str(suma_dann),
                     str(dt.today().day),
                     str(months[dt.today().month]), str(dt.today().year), str(punkt_4_6), str(suma_dann_dl),
                     str(punkt_7_8), str(inn_kpp1), str(ogrnip), str(leader_leasee_pod), str(imenyemoe), f7, b93]

        # создание ДЛ
        def replace_words_in_docx(docx_file, old_words, new_words):
            doc = Document(docx_file)

            for paragraph in doc.paragraphs:
                for i in range(len(old_words)):
                    if old_words[i] in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_words[i], str(new_words[i]))
                        # print(new_words[i])
                        # print(f'_____ {i=}')

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for i in range(len(old_words)):
                            if old_words[i] in cell.text:
                                cell.text = cell.text.replace(old_words[i], new_words[i])
            doc.save(dir_path / fr"ДЛ {inn_leasee}.docx")

        def change_font(docx_file, font_name):
            doc = Document(docx_file)

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(10)  # Установите желаемый размер шрифта

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = font_name
                                run.font.size = Pt(10)  # Установите желаемый размер шрифта

            doc.save(dir_path / fr"ДЛ {inn_leasee}.docx")

        if len(inn_leasee) == 12:
            if type_pl != 'on':
                replace_words_in_docx(dl_car_ip_path, old_words, new_words)
            else:
                replace_words_in_docx(dl_obor_ip_path, old_words,
                                      new_words)
        else:
            if type_pl != 'on':
                replace_words_in_docx(dl_car_ooo_path, old_words, new_words)
            else:
                replace_words_in_docx(dl_obor_ooo_path, old_words,
                                      new_words)

        change_font(dir_path / fr"ДЛ {inn_leasee}.docx", "Times New Roman")

    def grafic_punkty(inn_leasee, path_application, path_graphic, signatory, investor, currency_list, who_is_insure,
                      grafic):

        try:
            book = openpyxl.load_workbook(input_raschet_path, data_only=True)
            sheet = book[grafic]
            # print(f'Тут ГРАФИК {grafic}')
            suma_chislo = sheet['F7'].value
            replacements = {'{{ B9 }}': sheet['B9'].value.strftime('%d.%m.%Y'),
                            '{{ B10 }}': sheet['B10'].value.strftime('%d.%m.%Y'),
                            '{{ B11 }}': sheet['B11'].value.strftime('%d.%m.%Y'),
                            '{{ B12 }}': sheet['B12'].value.strftime('%d.%m.%Y'),
                            '{{ B13 }}': sheet['B13'].value.strftime('%d.%m.%Y'),
                            '{{ B14 }}': sheet['B14'].value.strftime('%d.%m.%Y'),
                            '{{ B15 }}': sheet['B15'].value.strftime('%d.%m.%Y'),
                            '{{ B16 }}': sheet['B16'].value.strftime('%d.%m.%Y'),
                            '{{ B17 }}': sheet['B17'].value.strftime('%d.%m.%Y'),
                            '{{ B18 }}': sheet['B18'].value.strftime('%d.%m.%Y'),
                            '{{ B19 }}': sheet['B19'].value.strftime('%d.%m.%Y'),
                            '{{ B20 }}': sheet['B20'].value.strftime('%d.%m.%Y'),
                            '{{ B21 }}': sheet['B21'].value.strftime('%d.%m.%Y'),
                            '{{ B22 }}': sheet['B22'].value.strftime('%d.%m.%Y'),
                            '{{ B23 }}': sheet['B23'].value.strftime('%d.%m.%Y'),
                            '{{ B24 }}': sheet['B24'].value.strftime('%d.%m.%Y'),
                            '{{ B25 }}': sheet['B25'].value.strftime('%d.%m.%Y'),
                            '{{ B26 }}': sheet['B26'].value.strftime('%d.%m.%Y'),
                            '{{ B27 }}': sheet['B27'].value.strftime('%d.%m.%Y'),
                            '{{ B28 }}': sheet['B28'].value.strftime('%d.%m.%Y'),
                            '{{ B29 }}': sheet['B29'].value.strftime('%d.%m.%Y'),
                            '{{ B30 }}': sheet['B30'].value.strftime('%d.%m.%Y'),
                            '{{ B31 }}': sheet['B31'].value.strftime('%d.%m.%Y'),
                            '{{ B32 }}': sheet['B32'].value.strftime('%d.%m.%Y'),
                            '{{ B33 }}': sheet['B33'].value.strftime('%d.%m.%Y'),
                            '{{ B34 }}': sheet['B34'].value.strftime('%d.%m.%Y'),
                            '{{ B35 }}': sheet['B35'].value.strftime('%d.%m.%Y'),
                            '{{ B36 }}': sheet['B36'].value.strftime('%d.%m.%Y'),
                            '{{ B37 }}': sheet['B37'].value.strftime('%d.%m.%Y'),
                            '{{ B38 }}': sheet['B38'].value.strftime('%d.%m.%Y'),
                            '{{ B39 }}': sheet['B39'].value.strftime('%d.%m.%Y'),
                            '{{ B40 }}': sheet['B40'].value.strftime('%d.%m.%Y'),
                            '{{ B41 }}': sheet['B41'].value.strftime('%d.%m.%Y'),
                            '{{ B42 }}': sheet['B42'].value.strftime('%d.%m.%Y'),
                            '{{ B43 }}': sheet['B43'].value.strftime('%d.%m.%Y'),
                            '{{ B44 }}': sheet['B44'].value.strftime('%d.%m.%Y'),
                            '{{ B45 }}': sheet['B45'].value.strftime('%d.%m.%Y'),
                            '{{ B46 }}': sheet['B46'].value.strftime('%d.%m.%Y'),
                            '{{ B47 }}': sheet['B47'].value.strftime('%d.%m.%Y'),
                            '{{ B48 }}': sheet['B48'].value.strftime('%d.%m.%Y'),
                            '{{ B49 }}': sheet['B49'].value.strftime('%d.%m.%Y'),
                            '{{ B50 }}': sheet['B50'].value.strftime('%d.%m.%Y'),
                            '{{ B51 }}': sheet['B51'].value.strftime('%d.%m.%Y'),
                            '{{ B52 }}': sheet['B52'].value.strftime('%d.%m.%Y'),
                            '{{ B53 }}': sheet['B53'].value.strftime('%d.%m.%Y'),
                            '{{ B54 }}': sheet['B54'].value.strftime('%d.%m.%Y'),
                            '{{ B55 }}': sheet['B55'].value.strftime('%d.%m.%Y'),
                            '{{ B56 }}': sheet['B56'].value.strftime('%d.%m.%Y'),
                            '{{ B57 }}': sheet['B57'].value.strftime('%d.%m.%Y'),
                            '{{ B58 }}': sheet['B58'].value.strftime('%d.%m.%Y'),
                            '{{ B59 }}': sheet['B59'].value.strftime('%d.%m.%Y'),
                            '{{ B60 }}': sheet['B60'].value.strftime('%d.%m.%Y'),
                            '{{ B61 }}': sheet['B61'].value.strftime('%d.%m.%Y'),
                            '{{ B62 }}': sheet['B62'].value.strftime('%d.%m.%Y'),
                            '{{ B63 }}': sheet['B63'].value.strftime('%d.%m.%Y'),
                            '{{ B64 }}': sheet['B64'].value.strftime('%d.%m.%Y'),
                            '{{ B65 }}': sheet['B65'].value.strftime('%d.%m.%Y'),
                            '{{ B66 }}': sheet['B66'].value.strftime('%d.%m.%Y'),
                            '{{ B67 }}': sheet['B67'].value.strftime('%d.%m.%Y'),
                            '{{ B68 }}': sheet['B68'].value.strftime('%d.%m.%Y'),
                            '{{ B93 }}': sheet['B93'].value.strftime('%d.%m.%Y'),
                            '{{ F7 }}': f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.', ',') if
                            f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F7'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                                    ',') + '0',
                            '{{ F8 }}': f"{round(float(sheet['F8'].value), 2):,}".replace(',', ' ').replace('.', ',') if
                            f"{round(float(sheet['F8'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F8'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                                    ',') + '0',
                            '{{ F9 }}': f"{round(float(sheet['F9'].value), 2):,}".replace(',', ' ').replace('.', ',') if
                            f"{round(float(sheet['F9'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F9'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                                    ',') + '0',
                            '{{ F10 }}': f"{round(float(sheet['F10'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F10'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F10'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F11 }}': f"{round(float(sheet['F11'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F11'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F11'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F12 }}': f"{round(float(sheet['F12'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F12'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F12'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F13 }}': f"{round(float(sheet['F13'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F13'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F13'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F14 }}': f"{round(float(sheet['F14'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F14'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F14'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F15 }}': f"{round(float(sheet['F15'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F15'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F15'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F16 }}': f"{round(float(sheet['F16'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F16'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F16'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F17 }}': f"{round(float(sheet['F17'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F17'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F17'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F18 }}': f"{round(float(sheet['F18'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F18'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F18'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F19 }}': f"{round(float(sheet['F19'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F19'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F19'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F20 }}': f"{round(float(sheet['F20'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F20'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F20'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F21 }}': f"{round(float(sheet['F21'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F21'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F21'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F22 }}': f"{round(float(sheet['F22'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F22'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F22'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F23 }}': f"{round(float(sheet['F23'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F23'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F23'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F24 }}': f"{round(float(sheet['F24'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F24'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F24'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F25 }}': f"{round(float(sheet['F25'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F25'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F25'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F26 }}': f"{round(float(sheet['F26'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F26'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F26'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F27 }}': f"{round(float(sheet['F27'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F27'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F27'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F28 }}': f"{round(float(sheet['F28'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F28'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F28'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F29 }}': f"{round(float(sheet['F29'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F29'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F29'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F30 }}': f"{round(float(sheet['F30'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F30'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F30'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F31 }}': f"{round(float(sheet['F31'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F31'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F31'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F32 }}': f"{round(float(sheet['F32'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F32'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F32'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F33 }}': f"{round(float(sheet['F33'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F33'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F33'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F34 }}': f"{round(float(sheet['F34'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F34'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F34'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F35 }}': f"{round(float(sheet['F35'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F35'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F35'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F36 }}': f"{round(float(sheet['F36'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F36'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F36'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F37 }}': f"{round(float(sheet['F37'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F37'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F37'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F38 }}': f"{round(float(sheet['F38'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F38'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F38'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F39 }}': f"{round(float(sheet['F39'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F39'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F39'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F40 }}': f"{round(float(sheet['F40'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F40'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F40'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F41 }}': f"{round(float(sheet['F41'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F41'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F41'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F42 }}': f"{round(float(sheet['F42'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F42'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F42'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F43 }}': f"{round(float(sheet['F43'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F43'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F43'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F44 }}': f"{round(float(sheet['F44'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F44'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F44'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F45 }}': f"{round(float(sheet['F45'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F45'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F45'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F46 }}': f"{round(float(sheet['F46'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F46'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F46'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F47 }}': f"{round(float(sheet['F47'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F47'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F47'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F48 }}': f"{round(float(sheet['F48'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F48'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F48'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F49 }}': f"{round(float(sheet['F49'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F49'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F49'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F50 }}': f"{round(float(sheet['F50'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F50'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F50'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F51 }}': f"{round(float(sheet['F51'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F51'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F51'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F52 }}': f"{round(float(sheet['F52'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F52'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F52'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F53 }}': f"{round(float(sheet['F53'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F53'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F53'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F54 }}': f"{round(float(sheet['F54'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F54'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F54'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F55 }}': f"{round(float(sheet['F55'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F55'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F55'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F56 }}': f"{round(float(sheet['F56'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F56'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F56'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F57 }}': f"{round(float(sheet['F57'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F57'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F57'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F58 }}': f"{round(float(sheet['F58'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F58'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F58'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F59 }}': f"{round(float(sheet['F59'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F59'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F59'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F60 }}': f"{round(float(sheet['F60'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F60'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F60'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F61 }}': f"{round(float(sheet['F61'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F61'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F61'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F62 }}': f"{round(float(sheet['F62'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F62'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F62'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F63 }}': f"{round(float(sheet['F63'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F63'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F63'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F64 }}': f"{round(float(sheet['F64'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F64'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F64'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F65 }}': f"{round(float(sheet['F65'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F65'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F65'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F66 }}': f"{round(float(sheet['F66'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F66'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F66'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F67 }}': f"{round(float(sheet['F67'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F67'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F67'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F68 }}': f"{round(float(sheet['F68'].value), 2):,}".replace(',', ' ').replace('.',
                                                                                                              ',') if
                            f"{round(float(sheet['F68'].value), 2):,}".replace(',', ' ').replace('.', ',')[
                                -3] == ',' else f"{round(float(sheet['F68'].value), 2):,}".replace(',', ' ').replace(
                                '.', ',') + '0',
                            '{{ F93 }}': f"{round(float(sheet['F93'].value), 2)}".replace('.',
                                                                                          ',') if f"{round(float(sheet['F93'].value), 2)}".replace(
                                '.', ',') == ',' else f"{round(float(sheet['F93'].value), 2)}".replace('.', ',') + '0',
                            }
            # logging.info(replacements)
            # print(replacements)
        except:
            replacements = {}
            # print('Не сработал график')

        doc = docx.Document(str(dir_path / fr"ДЛ {inn_leasee}.docx"))
        for para in doc.paragraphs:
            for old_word, new_word in replacements.items():
                for i, run in enumerate(para.runs):
                    if old_word in run.text:
                        # print(old_word, new_word)
                        run.text = run.text.replace(old_word, new_word)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_word, new_word in replacements.items():
                        for _ in cell.paragraphs:
                            for o, run in enumerate(_.runs):
                                # print(run.text)
                                if old_word in run.text:
                                    # print(old_word, new_word)
                                    run.text = run.text.replace(old_word, new_word)

        if vikup == '1000':
            for run in doc.paragraphs:
                if run.text == '4.6. Выкупная цена предмета лизинга, равная остаточной стоимости этого предмета лизинга, ' \
                               'рассчитывается в соответствии с действующим законодательством Российской Федерации с ' \
                               'учетом согласованной нормы амортизации. При этом выкупные платежи, уплаченные ' \
                               'Лизингополучателем в составе лизинговых платежей в соответствии с Приложением № 1 к ' \
                               'настоящему договору в качестве выкупных платежей, принимаются в зачет оплаты выкупной ' \
                               'цены предмета лизинга. В случае расторжения настоящего договора по вине ' \
                               'Лизингополучателя, в случае отказа от приемки предмета лизинга уплаченные ' \
                               'Лизингополучателем выкупные платежи возвращению не подлежат, а удерживаются ' \
                               'в качестве штрафа.':
                    doc.element.body.remove(run._element)

        for run in doc.paragraphs:
            if run.text == 'УДАЛИТЬ':
                doc.element.body.remove(run._element)

        if who_is_insure == 'ООО «ЛКМБ-РТ»':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'6.1. Лизингополучатель является страхователем в страховой компании, согласованной с Лизингодателем, по страхованию рисков утраты (хищения, угона), ущерба или полной гибели предмета лизинга. Выгодоприобретателем по договору страхования является {vigodo}. Согласованными с Лизингодателем страховыми компаниями являются страховые компании, соответствующие следующим требованиям:',
                    '- наличие действующей лицензии на осуществление страховой деятельности (в т.ч. требуемого вида страхования),',
                    '- осуществление фактической страховой деятельности без отзыва и приостановления лицензии не менее трех лет,',
                    '- имущество страховой компании не находится под арестом,',
                    '- отсутствие негативной информации, которая может поставить под сомнение способность страховой компании качественно и в срок выполнить принятые обязательства (например, уровень выплат по страховым договорам/полисам не менее 40%, и тп.),',
                    '- в отношении страховой компании не ведутся судебные процессы, способные оказать существенное негативное воздействие на ее деятельность, как ограничение, приостановление, либо отзыв лицензии на право осуществления страховой деятельности, вместе с тем, уменьшение стоимости активов страховой компании на 10% (Десять процентов) и более, размера чистой прибыли на 10% (Десять процентов) и более,',
                    '- категория по национальной (российской) рейтинговой шкале финансовой надежности не ниже ВВВ и размещением своих страховых резервов в соответствии с требованиями действующего законодательства Российской Федерации,',
                    '- условия предоставления страховой компанией страховых услуг в полной мере соответствуют требованиям действующего законодательства Российской Федерации,',
                    '- в отношении страховой организации не введена какая-либо из процедур банкротства согласно действующему законодательству, отсутствует решение государственного органа страхового надзора о приостановлении или отзыве лицензии на осуществление страховой деятельности;',
                    '- страховая организация имеет прибыль по результатам страховой деятельности по итогам годовой отчетности за последние два года, предшествующие году рассмотрения о признании страховой компании, отвечающей критериям лизинговой компании;',
                    '- наличие иных факторов, которые прямо или косвенно могут негативно влиять на финансово-экономическое состояние страховой компании;',
                    '- входит в перечень аккредитованных страховых компаний Инвестора.',
                    'Страховая сумма по страхованию на случай утраты (уничтожения, хищения) не может быть менее действительной стоимости предмета лизинга.',
                    'В целях получения разъяснений по условиям страхования Лизингополучатель имеет право обратиться к страховому брокеру Лизингодателя.',
                    f'6.1. Лизингодатель обязуется в течение 10 (Десяти) рабочих дней с даты передачи предмета лизинга по акту приема-передачи застраховать предмет лизинга в страховой компании, выбранной Лизингодателем, от хищения, угона, ущерба и полной гибели имущества на полную стоимость предмета лизинга с момента приема предмета лизинга на срок 12(двенадцать) месяцев. Выгодоприобретателем по договору страхования является {vigodo}.',
                    f'В последующий год и до окончания действия настоящего Договора Лизингополучатель является страхователем в страховой компании, согласованной с Лизингодателем, по страхованию рисков утраты (хищения, угона), ущерба или полной гибели предмета лизинга. Согласованными с Лизингодателем страховыми компаниями являются страховые компании, соответствующие следующим требованиям:']:
                    doc.element.body.remove(run._element)
        elif who_is_insure == '1 год ЛКМБ-РТ, дальше лизингополучатель':
            for run in doc.paragraphs:
                if run.text.strip() in ['6.1. Лизингодатель является страхователем в страховой компании, выбранной ' \
                                        'Лизингодателем, по страхованию рисков утраты (хищения, угона), ' \
                                        'ущерба или полной гибели предмета лизинга. Выгодоприобретателем по ' \
                                        f'договору страхования является {investor if investor != "ПАО «МКБ»" else "ООО «ЛКМБ-РТ»"}.',
                                        '6.1. Лизингодатель является страхователем в страховой компании, выбранной ' \
                                        'Лизингодателем, по страхованию рисков утраты (хищения, угона), ' \
                                        'ущерба или полной гибели предмета лизинга. Выгодоприобретателем по ' \
                                        f'договору страхования является Лизингодатель.',
                                        f'6.1. Лизингополучатель является страхователем в страховой компании, согласованной с Лизингодателем, по страхованию рисков утраты (хищения, угона), ущерба или полной гибели предмета лизинга. Выгодоприобретателем по договору страхования является {vigodo}. Согласованными с Лизингодателем страховыми компаниями являются страховые компании, соответствующие следующим требованиям:']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '6.1. Лизингодатель является страхователем в страховой компании, выбранной ' \
                    'Лизингодателем, по страхованию рисков утраты (хищения, угона), ' \
                    'ущерба или полной гибели предмета лизинга. Выгодоприобретателем по ' \
                    f'договору страхования является {investor if investor != "ПАО «МКБ»" else "ООО «ЛКМБ-РТ»"}.',
                    '6.1. Лизингодатель является страхователем в страховой компании, выбранной ' \
                    'Лизингодателем, по страхованию рисков утраты (хищения, угона), ' \
                    'ущерба или полной гибели предмета лизинга. Выгодоприобретателем по ' \
                    f'договору страхования является Лизингодатель.',
                    f'6.1. Лизингодатель обязуется в течение 10 (Десяти) рабочих дней с даты передачи предмета лизинга по акту приема-передачи застраховать предмет лизинга в страховой компании, выбранной Лизингодателем, от хищения, угона, ущерба и полной гибели имущества на полную стоимость предмета лизинга с момента приема предмета лизинга на срок 12(двенадцать) месяцев. Выгодоприобретателем по договору страхования является {vigodo}.',
                    f'В последующий год и до окончания действия настоящего Договора Лизингополучатель является страхователем в страховой компании, согласованной с Лизингодателем, по страхованию рисков утраты (хищения, угона), ущерба или полной гибели предмета лизинга. Согласованными с Лизингодателем страховыми компаниями являются страховые компании, соответствующие следующим требованиям:'
                ]:
                    doc.element.body.remove(run._element)

        if investor == 'ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»'.upper():
            for run in doc.paragraphs:
                if run.text.strip() in ['7.8. Лизингополучатель уведомлен:',
                                        '- о факте передаче предмета лизинга в залог Инвестору в счет исполнения обязательств Лизингодателя (далее в этом пункте также - Заемщик) по Кредитному договору;',
                                        '- о факте заключения договора залога прав требований Заемщика к Лизингополучателю по договору лизинга и передаче указанных прав требования Заемщика к Лизингополучателю в залог Инвестору в счет исполнения обязательств Заемщика по Кредитному договору;',
                                        '- о наличии права Инвестора как залогодержателя по договору залога прав требований Заемщика к Лизингополучателю по договору лизинга осуществить внесудебное обращение взыскания на права требования Заемщика к Лизингополучателю по договору лизинга.',
                                        'В связи с этим Лизингополучатель обязуется:',
                                        '- с даты получения уведомления АО «ПЕРВОУРАЛЬСКБАНК» как залогодержателя прав требований Заемщика к Лизингополучателю по договору лизинга осуществлять исполнения всех без исключения платежных обязательств Лизингополучателя по настоящему договору исключительно в адрес АО «ПЕРВОУРАЛЬСКБАНК» по реквизитам, которые будут указаны в соответствующем уведомлении.',
                                        '- уведомлять АО  «ПЕРВОУРАЛЬСКБАНК» о досрочном или частично досрочном гашении договора лизинга.',
                                        '- что предмет лизинга обременен залогом АО «Инвестторгбанк» по кредитному договору, в рамках которого осуществляется финансирование лизинговой сделки, и что права требования на получение платежей по договору лизинга переданы в залог Банку;',
                                        '- что Лизингополучатель при получении соответствующего уведомления от Банка (залогодержателя) обязан исполнять свое обязательство по договору лизинга Залогодержателю. При этом уведомление направляется любым из способов, определенных договором залога.']:
                    doc.element.body.remove(run._element)
        elif investor == 'АО «ПЕРВОУРАЛЬСКБАНК»'.upper():
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '7.8. Лизингополучатель уведомлен, что предмет залога и залог прав требования '
                    'по настоящему договору переданы в залог ПАО АКБ «МЕТАЛЛИНВЕСТБАНК».',
                    '- что предмет лизинга обременен залогом АО «Инвестторгбанк» по кредитному договору, в рамках которого осуществляется финансирование лизинговой сделки, и что права требования на получение платежей по договору лизинга переданы в залог Банку;',
                    '- что Лизингополучатель при получении соответствующего уведомления от Банка (залогодержателя) обязан исполнять свое обязательство по договору лизинга Залогодержателю. При этом уведомление направляется любым из способов, определенных договором залога.'
                ]:
                    doc.element.body.remove(run._element)
        elif investor == 'ИНВЕСТТОРГБАНК АО'.upper():
            for run in doc.paragraphs:
                if run.text.strip() in ['7.8. Лизингополучатель уведомлен:',
                                        '- о факте передаче предмета лизинга в залог Инвестору в счет исполнения обязательств Лизингодателя (далее в этом пункте также - Заемщик) по Кредитному договору;',
                                        '- о факте заключения договора залога прав требований Заемщика к Лизингополучателю по договору лизинга и передаче указанных прав требования Заемщика к Лизингополучателю в залог Инвестору в счет исполнения обязательств Заемщика по Кредитному договору;',
                                        '- о наличии права Инвестора как залогодержателя по договору залога прав требований Заемщика к Лизингополучателю по договору лизинга осуществить внесудебное обращение взыскания на права требования Заемщика к Лизингополучателю по договору лизинга.',
                                        'В связи с этим Лизингополучатель обязуется:',
                                        '- с даты получения уведомления АО «ПЕРВОУРАЛЬСКБАНК» как залогодержателя прав требований Заемщика к Лизингополучателю по договору лизинга осуществлять исполнения всех без исключения платежных обязательств Лизингополучателя по настоящему договору исключительно в адрес АО «ПЕРВОУРАЛЬСКБАНК» по реквизитам, которые будут указаны в соответствующем уведомлении.',
                                        '- уведомлять АО  «ПЕРВОУРАЛЬСКБАНК» о досрочном или частично досрочном гашении договора лизинга.',
                                        '7.8. Лизингополучатель уведомлен, что предмет залога и залог прав требования '
                                        'по настоящему договору переданы в залог ПАО АКБ «МЕТАЛЛИНВЕСТБАНК».'
                                        ]:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in ['7.8. Лизингополучатель уведомлен:',
                                        '- о факте передаче предмета лизинга в залог Инвестору в счет исполнения обязательств Лизингодателя (далее в этом пункте также - Заемщик) по Кредитному договору;',
                                        '- о факте заключения договора залога прав требований Заемщика к Лизингополучателю по договору лизинга и передаче указанных прав требования Заемщика к Лизингополучателю в залог Инвестору в счет исполнения обязательств Заемщика по Кредитному договору;',
                                        '- о наличии права Инвестора как залогодержателя по договору залога прав требований Заемщика к Лизингополучателю по договору лизинга осуществить внесудебное обращение взыскания на права требования Заемщика к Лизингополучателю по договору лизинга.',
                                        'В связи с этим Лизингополучатель обязуется:',
                                        '- с даты получения уведомления АО «ПЕРВОУРАЛЬСКБАНК» как залогодержателя прав требований Заемщика к Лизингополучателю по договору лизинга осуществлять исполнения всех без исключения платежных обязательств Лизингополучателя по настоящему договору исключительно в адрес АО «ПЕРВОУРАЛЬСКБАНК» по реквизитам, которые будут указаны в соответствующем уведомлении.',
                                        '- уведомлять АО  «ПЕРВОУРАЛЬСКБАНК» о досрочном или частично досрочном гашении договора лизинга.',
                                        '7.8. Лизингополучатель уведомлен, что предмет залога и залог прав требования '
                                        'по настоящему договору переданы в залог ПАО АКБ «МЕТАЛЛИНВЕСТБАНК».',
                                        '- что предмет лизинга обременен залогом АО «Инвестторгбанк» по кредитному договору, в рамках которого осуществляется финансирование лизинговой сделки, и что права требования на получение платежей по договору лизинга переданы в залог Банку;',
                                        '- что Лизингополучатель при получении соответствующего уведомления от Банка (залогодержателя) обязан исполнять свое обязательство по договору лизинга Залогодержателю. При этом уведомление направляется любым из способов, определенных договором залога.'
                                        ]:
                    doc.element.body.remove(run._element)

        doc.save(dir_path / fr"ДЛ {inn_leasee}.docx")

    logging.info(f"({current_user}) ЗАПУСК READ XLSX")

    read_xlsx(path_application, pl)  # читает эксель, после этого можно составлять ДЛ

    logging.info(f'({current_user}) ЗАПУСК CREATE DL')
    create_dl_dkp(inn_leasee, path_application, path_graphic, signatory, investor, currency_list, who_is_insure,
                  grafic, pl, number_dl, inn_seller)  # создается ДЛ
    grafic_punkty(inn_leasee, path_application, path_graphic, signatory, investor, currency_list, who_is_insure,
                  grafic)  # заполняет график платежей и убирает ненужные пункты
