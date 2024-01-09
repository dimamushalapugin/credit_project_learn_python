from docx import Document


def replace_words_in_bki(docx_file, old_words_bki, new_words_bki):
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        for i in range(len(old_words_bki)):
            if old_words_bki[i] in paragraph.text:
                paragraph.text = paragraph.text.replace(old_words_bki[i], str(new_words_bki[i]))

                # print(f'_____ {i=}')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i in range(len(old_words_bki)):
                    if old_words_bki[i] in cell.text:
                        cell.text = cell.text.replace(old_words_bki[i], str(new_words_bki[i]))
    new = new_words_bki[0].replace('"', '')
    if all(char.isdigit() for char in new_words_bki[1]):
        doc.save(fr'БКИ физ лицо {new}.docx')
    else:
        doc.save(fr'БКИ юр лицо {new_words_bki[1]}.docx')


def replace_bki(some1, some2, some3, some4, some5, some6, some7, some8, some9):
    print("Выполняю действия с юриком:", some1, some2, some3, some4, some5, some6, some7, some8, some9)
    date_string = some9
    year, month, day = date_string.split("-")
    months = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня',
              7: 'июля', 8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}

    old_words_bki = ["{{ fullname_company }}", "{{ ogrn_company }}", "{{ inn_company }}", "{{ address_company }}",
                     "{{ phone_company }}", "{{ leader_company }}", "{{ fio_company }}", "{{ doverka_company }}",
                     "{{ dt.today().day }}", "{{ months[dt.today().month] }}", "{{ dt.today().year }}"]

    new_words_bki = ([str(item) for item in [some2, some3, some1, some4, some5, some7, some6, some8]] +
                     [str(day), str(months[int(month)]), str(year)])
    print("old_words_bki:", old_words_bki, "new_words_bki:", new_words_bki)
    replace_words_in_bki(r"Согласие на получение кредитного отчета Юрлицо.docx", old_words_bki, new_words_bki)
    return old_words_bki, new_words_bki


def replace_bki_fiz(some1, some2, some3, some4, some5, some6, some7, some8, some9, some10, some11):
    print("Выполняю действия с физиком:", some1, some2, some3, some4, some5, some6, some7, some8, some9, some10, some11)
    date_mvd_string = some6
    year_mvd, month_mvd, day_mvd = date_mvd_string.split("-")
    datebirth_string = some9
    year_birth, month_birth, day_birth = datebirth_string.split("-")
    date_string = some11
    year, month, day = date_string.split("-")
    months = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня',
              7: 'июля', 8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}

    old_words_bki = ["{{ fio_fizik }}", "{{ ser_fizik }}", "{{ number_fizik }}", "{{ outputmvd_fizik }}",
                     "{{ daymvd_fizik }}", "{{ monthmvd_fizik }}", "{{ yearmvd_fizik }}", "{{ code_fizik }}",
                     "{{ birthplace_fizik }}", "{{ daybirthdate_fizik }}", "{{ monthbirthdate_fizik }}",
                     "{{ yearbirthdate_fizik }}", "{{ address_fizik }}", "{{ inn_fizik }}",
                     "{{ dt.today().day }}", "{{ months[dt.today().month] }}", "{{ dt.today().year }}"]

    new_words_bki = ([str(item) for item in [some2, some3, some4, some5, day_mvd, months[int(month_mvd)], year_mvd,
                                             some7, some8, day_birth, months[int(month_birth)], year_birth,
                                             some10, some1, str(day), str(months[int(month)]), str(year)]])
    print("old_words_bki:", old_words_bki, "new_words_bki:", new_words_bki)
    replace_words_in_bki(r"Согласие на получение кредитного отчета Физлицо.docx", old_words_bki, new_words_bki)
    return old_words_bki, new_words_bki
