import openpyxl
import os

from datetime import datetime as dt
from typing import Optional
from docx import Document
from num2words import num2words
from flask_login import current_user
from webapp.config import DADATA_BASE, dkp_car_ooo_path, dkp_obor_ooo_path
from webapp.risk.logger import logging
from webapp.risk.mongo_db import MongoDB
from pathlib import Path


def read_xlsx(path_application, pl):
    """
    :param path_application:
    :param pl:
    :return: 0: rekvizit_leasee_bik,
             1: rekvizit_leasee_cs_shet
             2: rekvizit_leasee_shet,
             3: rekvizit_leasee_bank,
             4: main_activity_leasee,
             5: fio_leader,
             6: email_leasee,
             7: phone_leasee,
             8: full_krakt_name_leasee,
             9: ustav_capital,
             10: date_regist,
             11: okpo_leasee,
             12: okato_leasee,
             13: ogrn_leasee,
             14: inn_seller_list,
             15: price_predmet_lizinga,
             16: predmet_lizinga,
             17: formatted_name_leader_leasee,
             18: leader_leasee,
             19: address_leasee_expluatazia,
             20: address_leasee,
             21: inn_kpp_leasee,
             22: full_name_leasee

    """
    # ip_or_kfh = 'Нет'
    wb = openpyxl.load_workbook(fr'{path_application}', data_only=True)
    # читаем страницу Заявление
    sheet_zayavlenie = wb['Заявление']
    full_name_leasee = sheet_zayavlenie['A5'].value
    inn_kpp_leasee = sheet_zayavlenie['D6'].value
    client_inn = inn_kpp_leasee.split('/')[0] if '/' in inn_kpp_leasee else inn_kpp_leasee

    formatted_name_leader_leasee = '-'
    leader_leasee = '-'
    address_leasee_expluatazia = '-'

    address_leasee = sheet_zayavlenie['H6'].value
    for number in range(24, sheet_zayavlenie.max_row + 2):
        if sheet_zayavlenie[
            f'B{number}'].value == (
                'Место эксплуатации предмета лизинга (для автотранспорта место стоянки/хранения) полный фактический адрес:'):
            address_leasee_expluatazia = sheet_zayavlenie[f'A{number + 1}'].value
        if sheet_zayavlenie[f'A{number}'].value == '(должность руководителя организации Заявителя)':
            leader_leasee = sheet_zayavlenie[f'A{number - 1}'].value
        if sheet_zayavlenie[f'O{number}'].value == '(расшифровка подписи)':
            formatted_name_leader_leasee = sheet_zayavlenie[f'O{number - 1}'].value

    predmet_lizinga = pl
    inn_seller_list = list(map(lambda x: sheet_zayavlenie[x].value, ['C11', 'C13', 'C15', 'C17']))
    for i in range(1, sheet_zayavlenie.max_row + 2):
        if sheet_zayavlenie[f'C{i}'].value == pl:
            price_predmet_lizinga = sheet_zayavlenie[f'P{i}'].value
            break
    else:
        price_predmet_lizinga = 0

    # читаем страницу Анкета Стр.1

    sheet_anketa_1_list = wb['Анкета_Стр.1']
    director_inn = sheet_anketa_1_list['H25'].value
    mongo = MongoDB(current_user)
    mongo.write_to_mongodb_bank_details(client_inn, sheet_anketa_1_list)
    mongo.write_to_mongodb_director_details(director_inn, sheet_anketa_1_list)
    ogrn_leasee = sheet_anketa_1_list['F7'].value
    okato_leasee = sheet_anketa_1_list['H7'].value
    okpo_leasee = sheet_anketa_1_list['J7'].value
    date_regist = sheet_anketa_1_list['E8'].value
    ustav_capital = sheet_anketa_1_list['J9'].value
    full_krakt_name_leasee = sheet_anketa_1_list[f'A6'].value

    phone_leasee = '-'
    email_leasee = '-'
    fio_leader = '-'
    main_activity_leasee = '-'
    rekvizit_leasee_bank = '-'
    rekvizit_leasee_shet = '-'
    rekvizit_leasee_cs_shet = '-'
    rekvizit_leasee_bik = '-'

    for number in range(8, sheet_anketa_1_list.max_row + 2):
        if sheet_anketa_1_list[f'A{number}'].value == '1.7         Телефон:':
            phone_leasee = sheet_anketa_1_list[f'C{number}'].value
        if sheet_anketa_1_list[f'E{number}'].value == '1.8 Эл. почта:':
            email_leasee = sheet_anketa_1_list[f'F{number}'].value
        if sheet_anketa_1_list[f'B{number}'].value == 'ФИО:':
            fio_leader = sheet_anketa_1_list[f'C{number}'].value
        if sheet_anketa_1_list[f'B{number}'].value == 'ОКВЭД с расшифровкой:':
            main_activity_leasee = sheet_anketa_1_list[f'E{number}'].value
        if sheet_anketa_1_list[f'F{number}'].value == 'Банк:':
            rekvizit_leasee_bank = sheet_anketa_1_list[f'G{number}'].value
        if sheet_anketa_1_list[f'A{number}'].value == 'Р/с:':
            rekvizit_leasee_shet = sheet_anketa_1_list[f'B{number}'].value
            rekvizit_leasee_cs_shet = sheet_anketa_1_list[f'F{number}'].value
            rekvizit_leasee_bik = sheet_anketa_1_list[f'I{number}'].value

    logging.info(f'({current_user}) Все данные успешно прочитаны')

    return (rekvizit_leasee_bik, rekvizit_leasee_cs_shet, rekvizit_leasee_shet, rekvizit_leasee_bank,
            main_activity_leasee, fio_leader, email_leasee, phone_leasee, full_krakt_name_leasee, ustav_capital,
            date_regist, okpo_leasee, okato_leasee, ogrn_leasee, inn_seller_list[0], price_predmet_lizinga,
            predmet_lizinga, formatted_name_leader_leasee, leader_leasee, address_leasee_expluatazia, address_leasee,
            inn_kpp_leasee, full_name_leasee)


def identification_lkmb_rt(signatory: str, investor: str):
    if signatory == 'Каюмов А. Д.':
        a_lkmb = 'Директор'
        lkmb_podpisant = 'Каюмов А. Д.'
        preambula_dolj_lkmb = 'Директора'
        preambula_fio_lkmb = 'Каюмова Айрата Дамировича'
        doverka_ustav_list = 'Устава'
        deystvuysh_list = 'действующего'
    elif signatory == 'Габдрахманов Р. Р.':
        a_lkmb = 'Заместитель директора'
        lkmb_podpisant = 'Габдрахманов Р. Р.'
        preambula_dolj_lkmb = 'Заместителя директора'
        preambula_fio_lkmb = 'Габдрахманова Рината Рафаэлевича'
        doverka_ustav_list = 'Доверенности от «17» марта 2020 года, удостоверенной Мальченковой  Евгенией Николаевной, нотариусом Казанского нотариального округа Республики Татарстан, зарегистрированной в реестре нотариальных действий за № 16/64-н/16-2020-7-317(бланк 16 АА 5665323)'
        deystvuysh_list = 'действующего'
    else:
        a_lkmb = 'Заместитель директора по финансам'
        lkmb_podpisant = 'Хасанова Д. Р.'
        preambula_dolj_lkmb = 'Заместителя директора по финансам'
        preambula_fio_lkmb = 'Хасановой Динары Ринатовны'
        doverka_ustav_list = 'Доверенности от «17» марта 2020 года, удостоверенной Мальченковой  Евгенией Николаевной, нотариусом Казанского нотариального округа Республики Татарстан, зарегистрированной в реестре нотариальных действий за № 16/64-н/16-2020-7-317(бланк 16 АА 5665323)'
        deystvuysh_list = 'действующей'

    if investor.upper() == 'ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»'.upper():
        r_chet_lkmb = '40701810000990000052'
        bank_rekv_lkmb = 'ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»'
        kor_chet_lkmb = '30101810300000000176'
        bik_lkmb = '044525176'
    elif investor.upper() == 'ПАО «МКБ»'.upper():
        r_chet_lkmb = '40701810900760000034'
        bank_rekv_lkmb = 'ПАО «МОСКОВСКИЙ КРЕДИТНЫЙ БАНК»'
        kor_chet_lkmb = '30101810745250000659'
        bik_lkmb = '044525659'
    elif investor.upper() == 'ИНВЕСТТОРГБАНК АО'.upper():
        r_chet_lkmb = '40701810071010300002'
        bank_rekv_lkmb = 'ИНВЕСТТОРГБАНК АО'
        kor_chet_lkmb = '30101810645250000267'
        bik_lkmb = '044525267'
    elif investor.upper() == 'АО «АЛЬФА-БАНК»'.upper():
        r_chet_lkmb = '40701810129930000005'
        bank_rekv_lkmb = 'ФИЛИАЛ «НИЖЕГОРОДСКИЙ» АО «АЛЬФА-БАНК»'
        kor_chet_lkmb = '30101810200000000824'
        bik_lkmb = '042202824'
    elif investor.upper() == 'АО «ПЕРВОУРАЛЬСКБАНК»'.upper():
        r_chet_lkmb = '40701810000010055037'
        bank_rekv_lkmb = 'АО «ПЕРВОУРАЛЬСКБАНК»'
        kor_chet_lkmb = '30101810565770000402'
        bik_lkmb = '046577402'
    elif investor.upper() == 'АО «СОЛИД БАНК»'.upper():
        r_chet_lkmb = '40701810105040011167'
        bank_rekv_lkmb = 'Московский филиал АО «СОЛИД БАНК»'
        kor_chet_lkmb = '30101810845250000795'
        bik_lkmb = '044525795'
    elif investor.upper() == 'АО КБ «УРАЛ ФД»'.upper():
        r_chet_lkmb = '40701810000000000962'
        bank_rekv_lkmb = 'АО КБ «УРАЛ ФД»'
        kor_chet_lkmb = '30101810800000000790'
        bik_lkmb = '045773790'
    elif investor.upper() == 'ПАО «Совкомбанк»'.upper():
        r_chet_lkmb = '40701810212020710089'
        bank_rekv_lkmb = 'ПАО «Совкомбанк»'
        kor_chet_lkmb = '30101810445250000360'
        bik_lkmb = '044525360'
    elif investor.upper() == 'ПАО «Сбербанк»'.upper():
        r_chet_lkmb = '40701810962000000214'
        bank_rekv_lkmb = 'ПАО «Сбербанк»'
        kor_chet_lkmb = '30101810900000000603'
        bik_lkmb = '042202603'
    else:
        r_chet_lkmb = '40702810100020002464'
        bank_rekv_lkmb = 'ПАО «АК БАРС» БАНК г. Казань'
        kor_chet_lkmb = '30101810000000000805'
        bik_lkmb = '049205805'
    return (a_lkmb, lkmb_podpisant, preambula_dolj_lkmb, preambula_fio_lkmb, doverka_ustav_list, deystvuysh_list,
            r_chet_lkmb, bank_rekv_lkmb, kor_chet_lkmb, bik_lkmb)


def indentification_pl(currency_list: str):
    currency_test = ''
    if currency_list == 'Рубль':
        currency_test = 'рублей'
    elif currency_list == 'Китайский юань':
        currency_test = 'юаней'
    elif currency_list == 'Доллар США':
        currency_test = 'долларов США'

    return currency_test


def number_to_words(suma_chislo123, curr):
    try:
        # Разбиваем строку на целую и десятичную часть
        suma_chislo123 = str(round(float(suma_chislo123), 2)).replace(',', '.') if \
            str(round(float(suma_chislo123), 2)).replace(',', '.')[-3] == '.' else str(
            round(float(suma_chislo123), 2)).replace(',', '.') + '0'
        parts = suma_chislo123.split(".")
        integer_part = parts[0]
        decimal_part = parts[1] if len(parts) > 1 else "00"

        # Преобразуем целую часть в число прописью
        integer_words = num2words(int(integer_part), lang='ru')

        # Определяем правильную форму для "рублей"
        if 10 < float(integer_part) % 100 < 20:
            valute_rub = "рублей"
        elif float(integer_part) % 10 == 1:
            valute_rub = "рубль"
        elif 1 < float(integer_part) % 10 < 5:
            valute_rub = "рубля"
        else:
            valute_rub = "рублей"

        # Преобразуем десятичную часть в число прописью
        if curr == 'Рубль':
            decimal_words = num2words(int(decimal_part), lang='ru', to='currency', currency='RUB')
        else:
            decimal_words = num2words(int(decimal_part), lang='ru')

        # Определяем правильную форму для "копеек"
        if 10 < float(decimal_part) % 100 < 20:
            valute_copeyka = "копеек"
        elif float(decimal_part) % 10 == 1:
            valute_copeyka = "копейка"
        elif 1 < float(decimal_part) % 10 <= 4:
            valute_copeyka = "копейки"
        else:
            valute_copeyka = "копеек"

        # Формируем итоговую строку
        if curr == 'Рубль':
            suma_dann1 = ((f"({integer_words}) {valute_rub} {decimal_part} {valute_copeyka}".strip()
                           .replace('ноль рублей', '').replace(',', '')))
        elif curr == 'Китайский юань':
            suma_dann1 = f"({integer_words}) китайских юаней {decimal_part} {valute_copeyka.replace('копеек', 'феней').replace('копейка', 'фень').replace('копейки', 'феня')}"
        elif curr == 'Доллар США':
            suma_dann1 = f"({integer_words}) долларов США {decimal_part} {valute_copeyka.replace('копеек', 'центов').replace('копейка', 'цент').replace('копейки', 'цента')}"
        elif curr == 'Евро':
            suma_dann1 = f"({integer_words}) евро {decimal_part} {valute_copeyka.replace('копеек', 'евроцентов').replace('копейка', 'евроцент').replace('копейки', 'евроцента')}"
        # print(suma_dann1)
        return suma_dann1
    except ValueError:
        return "Неверный формат числа"


def identification_leasee(leader_leasee):
    if leader_leasee.upper() == 'директор'.upper():
        leader_leasee_rod_padezh = 'Директора'
    elif leader_leasee.upper() == 'генеральный директор'.upper():
        leader_leasee_rod_padezh = 'Генерального директора'
    elif leader_leasee.upper() == 'исполняющий обязанности директора'.upper():
        leader_leasee_rod_padezh = 'ИО директора'
    elif leader_leasee.upper() == 'управляющий'.upper():
        leader_leasee_rod_padezh = 'Управляющего'
    elif leader_leasee.upper() == 'председатель'.upper():
        leader_leasee_rod_padezh = 'Председателя'
    elif leader_leasee.upper() == 'президент'.upper():
        leader_leasee_rod_padezh = 'Президента'
    elif leader_leasee.upper() == 'вице-президент'.upper():
        leader_leasee_rod_padezh = 'Вице-Президента'
    elif leader_leasee.upper() == 'исполнительный директор'.upper():
        leader_leasee_rod_padezh = 'Исполнительного Директора'
    elif leader_leasee.upper() == 'финансовый директор'.upper():
        leader_leasee_rod_padezh = 'Финансового Директора'
    elif leader_leasee.upper() == 'коммерческий директор'.upper():
        leader_leasee_rod_padezh = 'Коммерческого Директора'
    elif leader_leasee.upper() == 'глава'.upper():
        leader_leasee_rod_padezh = 'Главы'
    elif leader_leasee.upper() == 'ректор'.upper():
        leader_leasee_rod_padezh = 'Ректора'
    else:
        leader_leasee_rod_padezh = ''
    leader_leasee_pod = leader_leasee
    return leader_leasee_rod_padezh, leader_leasee_pod


def start_filling_agreement_dkp(path_application: str, inn_client: str, inn_seller: str, numb_dl_dkp: str,
                                signatory: str, investor: str, currency: str, pl: str, equipment_or_not: Optional[str],
                                pl_new_or_not: str, payment_order: str, place: str, acts: str, diadok: str,
                                pnr: Optional[str], house: Optional[str], learn: Optional[str], stock: str):
    """
    Аргументы:
    1. path_application: Путь до заявки. | Не может быть None
    2. inn_client: ИНН лизингополучателя. | Пользователь вручную вводит ИНН | Вывод: "1234567899"
                   | Не может быть None
    3. inn_seller: ИНН продавца | Пользователь вручную вводит ИНН | Вывод: "1234567899"
                   | Не может быть None
    4. numb_dl_dkp: Номер ДЛ/ДКП | Пользователь вручную вводит номер ДЛ/ДКП |Вывод: "1/23". | Не может быть None
    5. signatory: Подписант | values = ['Каюмов А. Д.', 'Габдрахманов Р. Р.', 'Хасанова Д. Р.']. |
                  В зависимости от выбора, может быть одним из values. Вывод: "Хасанова Д. Р."
                  | Не может быть None
    6. investor: Инвестор | values = ["ПАО «АК БАРС» БАНК", "АО «АЛЬФА-БАНК»", "ПАО АКБ «МЕТАЛЛИНВЕСТБАНК»",
                 "ПАО «МКБ»", "АО «ПЕРВОУРАЛЬСКБАНК»", "АО «СОЛИД БАНК»", "АО КБ «УРАЛ ФД»", "ИНВЕСТТОРГБАНК АО",
                 "ООО «ЛКМБ-РТ»"]. | В зависимости от выбора, может быть одним из values. Вывод: "ПАО «АК БАРС» БАНК"
                 | Не может быть None
    7. currency: Валюта указанная в ДЛ/ДКП | values = ["Рубль", "Китайский юань", "Доллар США"]. |
                 В зависимости от выбора, может быть одним из values. Вывод: "Рубль"
                 | Не может быть None
    8. pl: Предмет лизинга | values = ["0", "1", "2", "3"]. |
           В зависимости от выбора, может быть один из values. Вывод: "0"
           | Не может быть None
    9. equipment_or_not: Оборудование (чек-лист) | values = ["on", None]. |
                  В зависимости от выбора, может быть одним из values. Вывод: "on"
                  | Может быть None, если чек-лист не прожат
    10. pl_new_or_not: Предмет лизинга Б/У или Новый | values = ["used","used_with_garantee", "new"]. |
                       В зависимости от выбора, может быть один из values. Вывод: "used"
                       | Не может быть None
    11. payment_order: Порядок оплаты поставщику
                       | Пользователь вручную вводит платежи (в %) через пробел, например: 20 80)
                       | Вывод: "5 15 80" | Не может быть None
    12. place: Место отгрузки ПЛ (По месту рег. продавца или По месту рег. ЛП) | values = ["продавец", "лп"]. |
               В зависимости от выбора, может быть один из values. Вывод: "лп"
               | Не может быть None
    13. acts: Подписание актов приема-передачи (В момент отгрузки со склада продавца или По месту эксплуатации)
              | values = ["склад", "эксплуатация"]. | В зависимости от выбора, может быть один из values. Вывод: "склад"
              | Не может быть None
    14. diadok: Способ подписания договора (Диадок, ЭЦП или Живая подпись)
              | values = ["диадок", "эцп", "живая"] | В зависимости от выбора, может быть один из values. Вывод: "живая"
              | Не может быть None
    15. stock: Способ оплаты поставщику (Обычный в рублях, ЦБ, Биржа)
              | values = ["ЦБ", "Биржа"] | В зависимости от выбора, может быть один из values. Вывод: "ЦБ"
              | Может быть None, если выбрана валюта "руб."
    ====================================================================================================================
    Далее идут 3 последних аргумента. Они появляются у пользователя, если прожат чек-бокс на оборудование
    (equipment_or_not == 'on'). Если чек-бокс не нажат (equipment_or_not == None), тогда все три аргумента будут None
    ====================================================================================================================
    16. pnr: Необходимость пуско-наладочных работ | values = ["Да", "Нет"]
             | В зависимости от выбора, может быть один из values. Вывод: "Да"
             | Может быть None | Если None значит ПЛ не оборудование
    17. house: Необходимость подготовки помещения | values = ["Да", "Нет"]
               | В зависимости от выбора, может быть один из values. Вывод: "Нет"
               | Может быть None | Если None значит ПЛ не оборудование
    18. learn: Необходимость обучения персонала | values = ["Да", "Нет"]
               | В зависимости от выбора, может быть один из values. Вывод: "Нет"
               | Может быть None | Если None значит ПЛ не оборудование
    """

    #  TODO: Валидатор добавить ошибки
    def quantity_payment_order():
        summ = 0
        for num in payment_order.split(' '):
            summ += float(num)
        # print(summ)

        if summ == 100:
            print("Суммы равны 100")
        else:
            print(
                "Суммы не равны 100, надо написать пользователю что он даун и не умеет считать! Flash должен выскочить")

    # 1
    def equipment_valute():
        if pl_new_or_not == 'new':
            new_old_pl = 'новое'
        else:
            new_old_pl = 'бывшее в употреблении'

        if investor == 'АО «ПЕРВОУРАЛЬСКБАНК»':
            pb_vizor = (', и для АО «Первоуральскбанк» - фотографии, '
                        'идентифицирующие предмет лизинга (не менее 5 шт.), используя программный комплекс PB-visor')
        else:
            pb_vizor = ''

        if equipment_or_not == None:
            identif_punkt_3_1_1 = ('Предмету лизинга присваивается идентификационный номер, '
                                   'который  вместе с другими данными выбивается на металлическом шильдике (табличке), '
                                   'расположенном на предмете лизинга на видном месте. '
                                   'Этот номер вносится и в паспорт предмета лизинга.')
            identif_punkt_3_1_3 = 'комплектом ключей зажигания и '
            punkt_3_1_9 = '3.1.9.'
            punkt_3_3_3_key = 'и комплекта ключей зажигания '
            punkt_3_3_3_key2 = 'и второй комплект ключей '
            punkt_3_1_6 = ', 3.1.6 '
            punkt_3_3_7 = ',3.3.4'
            punkt_8_2 = 'четырех (четвертый – для регистрирующего органа)'
            pril_1_2 = 'ключи зажигания -1 компл., '
            pril_1_3 = 'ключи зажигания – 1 компл., подлинник ПТС /ПСМ,'
        else:
            identif_punkt_3_1_1 = ''
            identif_punkt_3_1_3 = ''
            punkt_3_1_9 = '3.1.9.'
            punkt_3_3_3_key = ''
            punkt_3_3_3_key2 = ''
            punkt_3_1_6 = ' '
            punkt_3_3_7 = '3.3.7'
            punkt_8_2 = 'трех'
            pril_1_2 = ''
            pril_1_3 = ''

        if equipment_or_not == 'on' and pnr == 'Нет':
            punkt_4_7_1 = ''
            punkt_5_3 = ''
        if equipment_or_not == 'on' and pnr == 'Да':
            punkt_4_7_1 = 'из полученного от Лизингополучателя первоначального платежа '
            punkt_5_3 = ' акта приема пуско-наладочных работ '
            punkt_6_4 = ', 3.3.5, 3.3.6  '
        else:
            punkt_4_7_1 = ''
            punkt_5_3 = ''
            punkt_6_4 = ''

        if investor == 'ПАО «АК БАРС» БАНК':
            punkt_8_3 = ('Вместе с тем, Продавец и Лизингополучатель дают согласие '
                         'Покупателю на передачу в залог обязательственных прав с '
                         'последующим залогом имущества по настоящему Договору.')
        else:
            punkt_8_3 = ''

        if stock != 'ЦБ' or stock != 'Биржа':
            kurs = 'Обычная оплата'  # еще возможна МОсбиржа  или валюта по ЦБ
        elif stock == 'Биржа':
            kurs = 'moex'
        else:
            kurs = 'ЦБ'
        type_currency = ''
        equivalent_currency = ''
        if currency == 'Рубль':
            type_currency = 'российских рублях'
        elif currency == 'китайских юанях':
            type_currency = 'китайских юанях'
            equivalent_currency = 'в рублях, эквивалентную'
        elif currency == 'Доллар США':
            type_currency = 'долларах США'
            equivalent_currency = 'в рублях, эквивалентную'
        elif currency == 'евро':
            type_currency = 'евро'
            equivalent_currency = 'в рублях, эквивалентную'
        return (new_old_pl, pb_vizor, identif_punkt_3_1_1, identif_punkt_3_1_3, punkt_3_1_9, punkt_3_3_3_key,
                punkt_3_3_3_key2, punkt_3_1_6, punkt_3_3_7, punkt_8_2, pril_1_2, pril_1_3, punkt_4_7_1, punkt_5_3,
                punkt_6_4, punkt_8_3, type_currency, equivalent_currency, kurs)

    def percent_to_word(number: str):
        try:
            # Разбиваем строку на целую и десятичную часть
            suma_chislo = str(round(float(number), 2)).replace(',', '.') if \
                str(round(float(number), 2)).replace(',', '.')[-3] == '.' else str(
                round(float(number), 2)).replace(',', '.') + '0'
            parts = suma_chislo.split(".")
            integer_part = parts[0]
            decimal_part = parts[1] if len(parts) > 1 else "00"

            # Преобразуем целую часть в число прописью
            integer_words = num2words(int(integer_part), lang='ru')
            # print(integer_part)

            # Преобразуем десятичную часть в число прописью
            decimal_words = num2words(int(decimal_part), lang='ru')
            # print(decimal_part)

            # формируем итоговую строку
            if int(integer_part) in [1, 21, 31, 41, 51, 61, 71, 81, 91] and decimal_part == '00':
                punkt_2_3_1_numb_pr12345 = f'{integer_words} процент'
            elif int(integer_part) in [5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 25, 26, 27, 28, 29,
                                       30, 35, 36, 37, 38, 38, 40, 45, 46, 47, 48, 49, 50, 55, 56, 57, 58, 59, 60, 65,
                                       66, 67, 68, 69, 70, 75, 76, 77, 78, 79, 80, 85, 86, 87, 88, 89, 90, 95, 96, 97,
                                       98, 99] and decimal_part == '00':
                punkt_2_3_1_numb_pr12345 = f'{integer_words} процентов'
            elif int(integer_part) in [2, 3, 4, 22, 23, 24, 32, 33, 34, 42, 43, 44, 52, 53, 54, 62, 63, 64, 72, 73, 74,
                                       82, 83, 84, 92, 93, 94] and decimal_part == '00':
                punkt_2_3_1_numb_pr12345 = f'{integer_words} процента'
            else:
                punkt_2_3_1_numb_pr12345 = f'{integer_words} целых {decimal_words} сотых процентов'
            return punkt_2_3_1_numb_pr12345
        except ValueError:
            return "Неверный формат числа"

    def payment_for_dkp(price_entry):
        # print(payment_order)
        punkt_2_3_1_pay = ''
        punkt_2_3_2_pay = ''
        punkt_2_3_3_pay = ''
        punkt_2_3_4_pay = ''
        punkt_2_3_5_pay = ''
        punkt_2_3_1_num = ''
        punkt_2_3_2_num = ''
        punkt_2_3_3_num = ''
        punkt_2_3_4_num = ''
        punkt_2_3_5_num = ''
        punkt_2_3_1_numb_pr = ''
        punkt_2_3_2_numb_pr = ''
        punkt_2_3_3_numb_pr = ''
        punkt_2_3_4_numb_pr = ''
        punkt_2_3_5_numb_pr = ''
        payment_1 = ''
        payment_2 = ''
        payment_3 = ''
        payment_4 = ''
        payment_5 = ''
        payment_1_propis = ''
        payment_2_propis = ''
        payment_3_propis = ''
        payment_4_propis = ''
        payment_5_propis = ''
        "информация по платежам, пункт 2.3"
        if len(payment_order.split(' ')) == 1:
            punkt_2_3_1_pay = 'Окончательный'
            punkt_2_3_1_num = payment_order.split(' ')[0]
            punkt_2_3_1_numb_pr = percent_to_word(str(punkt_2_3_1_num))
            payment_1 = round((float(price_entry) * float(payment_order.split(' ')[0])) / 100, 2)
            payment_1_propis = number_to_words(payment_1, currency)
            payment_1 = payment_1 if str(payment_1)[-3] == '.' else str(payment_1) + '0'
        elif len(payment_order.split(' ')) == 2:
            punkt_2_3_1_pay = 'Первый'
            punkt_2_3_2_pay = 'Окончательный'
            punkt_2_3_1_num = payment_order.split(' ')[0]
            punkt_2_3_2_num = payment_order.split(' ')[1]
            punkt_2_3_1_numb_pr = percent_to_word(str(punkt_2_3_1_num))
            punkt_2_3_2_numb_pr = percent_to_word(str(punkt_2_3_2_num))
            payment_1 = round((float(price_entry) * float(payment_order.split(' ')[0])) / 100, 2)
            payment_2 = round((float(price_entry) * float(payment_order.split(' ')[1])) / 100, 2)
            payment_1_propis = number_to_words(payment_1, currency)
            payment_2_propis = number_to_words(payment_2, currency)
            payment_1 = payment_1 if str(payment_1)[-3] == '.' else str(payment_1) + '0'
            payment_2 = payment_2 if str(payment_2)[-3] == '.' else str(payment_2) + '0'
        elif len(payment_order.split(' ')) == 3:
            punkt_2_3_1_pay = 'Первый'
            punkt_2_3_2_pay = 'Второй'
            punkt_2_3_3_pay = 'Окончательный'
            punkt_2_3_1_num = payment_order.split(' ')[0]
            punkt_2_3_2_num = payment_order.split(' ')[1]
            punkt_2_3_3_num = payment_order.split(' ')[2]
            punkt_2_3_1_numb_pr = percent_to_word(str(punkt_2_3_1_num))
            punkt_2_3_2_numb_pr = percent_to_word(str(punkt_2_3_2_num))
            punkt_2_3_3_numb_pr = percent_to_word(str(punkt_2_3_3_num))
            payment_1 = round((float(price_entry) * float(payment_order.split(' ')[0])) / 100, 2)
            payment_2 = round((float(price_entry) * float(payment_order.split(' ')[1])) / 100, 2)
            payment_3 = round((float(price_entry) * float(payment_order.split(' ')[2])) / 100, 2)
            payment_1_propis = number_to_words(payment_1, currency)
            payment_2_propis = number_to_words(payment_2, currency)
            payment_3_propis = number_to_words(payment_3, currency)
            payment_1 = payment_1 if str(payment_1)[-3] == '.' else str(payment_1) + '0'
            payment_2 = payment_2 if str(payment_2)[-3] == '.' else str(payment_2) + '0'
            payment_3 = payment_3 if str(payment_3)[-3] == '.' else str(payment_3) + '0'
        elif len(payment_order.split(' ')) == 4:
            punkt_2_3_1_pay = 'Первый'
            punkt_2_3_2_pay = 'Второй'
            punkt_2_3_3_pay = 'Третий'
            punkt_2_3_4_pay = 'Окончательный'
            punkt_2_3_1_num = payment_order.split(' ')[0]
            punkt_2_3_2_num = payment_order.split(' ')[1]
            punkt_2_3_3_num = payment_order.split(' ')[2]
            punkt_2_3_4_num = payment_order.split(' ')[3]
            punkt_2_3_1_numb_pr = percent_to_word(str(punkt_2_3_1_num))
            punkt_2_3_2_numb_pr = percent_to_word(str(punkt_2_3_2_num))
            punkt_2_3_3_numb_pr = percent_to_word(str(punkt_2_3_3_num))
            punkt_2_3_4_numb_pr = percent_to_word(str(punkt_2_3_4_num))
            payment_1 = round((float(price_entry) * float(payment_order.split(' ')[0])) / 100, 2)
            payment_2 = round((float(price_entry) * float(payment_order.split(' ')[1])) / 100, 2)
            payment_3 = round((float(price_entry) * float(payment_order.split(' ')[2])) / 100, 2)
            payment_4 = round((float(price_entry) * float(payment_order.split(' ')[3])) / 100, 2)
            payment_1_propis = number_to_words(payment_1, currency)
            payment_2_propis = number_to_words(payment_2, currency)
            payment_3_propis = number_to_words(payment_3, currency)
            payment_4_propis = number_to_words(payment_4, currency)
            payment_1 = payment_1 if str(payment_1)[-3] == '.' else str(payment_1) + '0'
            payment_2 = payment_2 if str(payment_2)[-3] == '.' else str(payment_2) + '0'
            payment_3 = payment_3 if str(payment_3)[-3] == '.' else str(payment_3) + '0'
            payment_4 = payment_4 if str(payment_4)[-3] == '.' else str(payment_4) + '0'
        elif len(payment_order.split(' ')) == 5:
            punkt_2_3_1_pay = 'Первый'
            punkt_2_3_2_pay = 'Второй'
            punkt_2_3_3_pay = 'Третий'
            punkt_2_3_4_pay = 'Четвертый'
            punkt_2_3_5_pay = 'Окончательный'
            punkt_2_3_1_num = payment_order.split(' ')[0]
            punkt_2_3_2_num = payment_order.split(' ')[1]
            punkt_2_3_3_num = payment_order.split(' ')[2]
            punkt_2_3_4_num = payment_order.split(' ')[3]
            punkt_2_3_5_num = payment_order.split(' ')[4]
            punkt_2_3_1_numb_pr = percent_to_word(str(punkt_2_3_1_num))
            punkt_2_3_2_numb_pr = percent_to_word(str(punkt_2_3_2_num))
            punkt_2_3_3_numb_pr = percent_to_word(str(punkt_2_3_3_num))
            punkt_2_3_4_numb_pr = percent_to_word(str(punkt_2_3_4_num))
            punkt_2_3_5_numb_pr = percent_to_word(str(punkt_2_3_5_num))
            payment_1 = round((float(price_entry) * float(payment_order.split(' ')[0])) / 100, 2)
            payment_2 = round((float(price_entry) * float(payment_order.split(' ')[1])) / 100, 2)
            payment_3 = round((float(price_entry) * float(payment_order.split(' ')[2])) / 100, 2)
            payment_4 = round((float(price_entry) * float(payment_order.split(' ')[3])) / 100, 2)
            payment_5 = round((float(price_entry) * float(payment_order.split(' ')[4])) / 100, 2)
            payment_1_propis = number_to_words(payment_1, currency)
            payment_2_propis = number_to_words(payment_2, currency)
            payment_3_propis = number_to_words(payment_3, currency)
            payment_4_propis = number_to_words(payment_4, currency)
            payment_5_propis = number_to_words(payment_5, currency)
            payment_1 = payment_1 if str(payment_1)[-3] == '.' else str(payment_1) + '0'
            payment_2 = payment_2 if str(payment_2)[-3] == '.' else str(payment_2) + '0'
            payment_3 = payment_3 if str(payment_3)[-3] == '.' else str(payment_3) + '0'
            payment_4 = payment_4 if str(payment_4)[-3] == '.' else str(payment_4) + '0'
            payment_5 = payment_5 if str(payment_5)[-3] == '.' else str(payment_5) + '0'
        return [punkt_2_3_1_pay, punkt_2_3_2_pay, punkt_2_3_3_pay, punkt_2_3_4_pay, punkt_2_3_5_pay, punkt_2_3_1_num,
                punkt_2_3_2_num, punkt_2_3_3_num, punkt_2_3_4_num, punkt_2_3_5_num, punkt_2_3_1_numb_pr,
                punkt_2_3_2_numb_pr, punkt_2_3_3_numb_pr, punkt_2_3_4_numb_pr, punkt_2_3_5_numb_pr, payment_1,
                payment_2, payment_3, payment_4, payment_5, payment_1_propis, payment_2_propis, payment_3_propis,
                payment_4_propis, payment_5_propis]

    def result_dadata():
        result_dkp = DADATA_BASE.find_by_id("party", inn_seller)
        return result_dkp

    def some_info_seller(result_dkp):
        leader_seller = ''
        try:
            leader_seller = result_dkp[0]['data']['management']['post']
            # print(leader_seller)
        except:
            leader_seller = result_dkp[0]['data']['opf']['short']
            # print(leader_seller)
        print('43535')
        print(leader_seller)
        if leader_seller.upper() == 'директор'.upper():
            leader_seller_rod_padezh = 'Директора'
        elif leader_seller.upper() == 'генеральный директор'.upper():
            leader_seller_rod_padezh = 'Генерального директора'
        elif leader_seller.upper() == 'исполняющий обязанности директора'.upper():
            leader_seller_rod_padezh = 'ИО директора'
        elif leader_seller.upper() == 'управляющий'.upper():
            leader_seller_rod_padezh = 'Управляющего'
        elif leader_seller.upper() == 'председатель'.upper():
            leader_seller_rod_padezh = 'Председателя'
        elif leader_seller.upper() == 'президент'.upper():
            leader_seller_rod_padezh = 'Президента'
        elif leader_seller.upper() == 'вице-президент'.upper():
            leader_seller_rod_padezh = 'Вице-Президента'
        elif leader_seller.upper() == 'исполнительный директор'.upper():
            leader_seller_rod_padezh = 'Исполнительного Директора'
        elif leader_seller.upper() == 'финансовый директор'.upper():
            leader_seller_rod_padezh = 'Финансового Директора'
        elif leader_seller.upper() == 'коммерческий директор'.upper():
            leader_seller_rod_padezh = 'Коммерческого Директора'
        elif leader_seller.upper() == 'глава'.upper():
            leader_seller_rod_padezh = 'Главы'
        elif leader_seller.upper() == 'глава кфх'.upper():
            leader_seller_rod_padezh = 'Главы'
        elif leader_seller.upper() == 'ректор'.upper():
            leader_seller_rod_padezh = 'Ректора'
        else:
            leader_seller_rod_padezh = ''
        leader_seller_pod = leader_seller
        print('Leaderseller123')
        print(leader_seller, leader_seller_rod_padezh)
        return leader_seller, leader_seller_rod_padezh, leader_seller_pod

    def full_rekviti_seller(result_dkp):
        ip_or_kfh_dkp = 'Нет'
        if result_dkp[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
            ip_or_kfh_dkp = 'Да'
            logging.info('12131')
            logging.info(ip_or_kfh_dkp)
        if ip_or_kfh_dkp == 'Нет':
            inn_kpp_seller = inn_seller + '/' + result_dkp[0]['data']['kpp']
            leader_seller = result_dkp[0]['data']['management']['post']
            fio_leader_seller = result_dkp[0]['data']['management']['name']
            logging.info('123')
            logging.info(ip_or_kfh_dkp)
        else:
            inn_kpp_seller = inn_seller
            type_business_dkp = result_dkp[0]['data']['opf']['short']
            leader_seller = result_dkp[0]['data']['opf']['short']
            if result_dkp[0]['data']['opf']['short'] == 'ИП':
                leader_seller = 'Индивидуальный предприниматель'
            else:
                leader_seller = 'Директор'
            fio_leader_seller = result_dkp[0]['data']['name']['full']
        last_name_dkp, first_name_dkp, patronymic_name_dkp = fio_leader_seller.split()

        # Get the initial of the first name
        first_name_initial_dkp = first_name_dkp[0]

        # Get the initial of the patronymic name
        patronymic_initial_dkp = patronymic_name_dkp[0]

        # Combine the last name and initials in the desired format
        formatted_name_leader_seller = f'{last_name_dkp} {first_name_initial_dkp}.{patronymic_initial_dkp}.'
        address_seller_dkp = result_dkp[0]['data']['address']['unrestricted_value']
        ogrn_seller = result_dkp[0]['data']['ogrn']
        return (ip_or_kfh_dkp, inn_kpp_seller, ogrn_seller, leader_seller, formatted_name_leader_seller,
                fio_leader_seller, address_seller_dkp)

    def seller_dkp_all():
        result_dkp = DADATA_BASE.find_by_id("party", inn_seller)
        full_name_seller = result_dkp[0]['data']['name']['full_with_opf']
        # print(full_name_seller)
        doverka_ustav_seller = 'Устава'
        for elem in full_name_seller.split():
            if elem in ['Индивидуальный', 'предприниматель', 'хозяйства']:
                doverka_ustav_seller = f'Свидетельства о государственной регистрации физического лица в качестве индивидуального предпринимателя ОГРНИП '
        return full_name_seller, doverka_ustav_seller

    def deistv_seller(result_dkp, fio, gender):
        imenyemoe_dkp = 'именуемое'
        try:
            if gender == 'М':
                deystvuysh_list_seller = 'действующего'
                if result_dkp[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_seller = 'действующий'
                imenyemoe_dkp = 'именуемый'
            else:
                deystvuysh_list_seller = 'действующая'
                imenyemoe_dkp = 'именуемая'
        except:
            try:
                if result_dkp[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_seller = 'действующая'
                else:
                    deystvuysh_list_seller = '-'
            except:
                imenyemoe_dkp = 'именуемая'
        return deystvuysh_list_seller, imenyemoe_dkp

    inn_leasee1 = read_xlsx(path_application, pl)

    def result_dadata_leasee():
        result_dkp_leasee = DADATA_BASE.find_by_id("party", inn_leasee1)
        if result_dkp_leasee[0]['data']['opf']['short'] in ['ИП', 'ГКФХ', 'КФХ']:
            inn_kpp1 = 'ИНН'
            put_padezh_podpisant_rg = ''
        else:
            imenyemoe = 'именуемое'
        return result_dkp_leasee

    def rod_padezh_fio_leader(fio):
        # dadata = Dadata(DADATA_TOKEN, DADATA_SECRET)
        logging.info(f"({fio})")
        put_padezh_podpisant = DADATA_BASE.clean("name", fio)
        # print(put_padezh_podpisant)
        return put_padezh_podpisant

    rod_padezh_fio_leader = rod_padezh_fio_leader(read_xlsx(path_application, pl)[5])
    ip_or_not = read_xlsx(path_application, pl)

    def addicted_info_leasee(date_regist, ogrn_leasee):
        deystvuysh_list_leasee = 'действующей'
        imenyemoe = 'именуемое'
        try:
            put_padezh_podpisant_rg = rod_padezh_fio_leader['result_genitive']
        except:
            put_padezh_podpisant_rg = ''
        # print(f'123 {put_padezh_podpisant_rg}')
        doverka_ustav_leasee = 'Устава'
        for elem in ip_or_not[22].split():
            if elem in ['Индивидуальный', 'предприниматель', 'хозяйства']:
                doverka_ustav_leasee = f'Свидетельства о государственной регистрации физического лица в качестве индивидуального предпринимателя от {date_regist}, ОГРНИП {ogrn_leasee}'

        try:
            if rod_padezh_fio_leader['gender'] == 'М':
                deystvuysh_list_leasee = 'действующего'
                if result_dadata_leasee()[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_leasee = 'действующий'
                imenyemoe = 'именуемый'
        except:
            try:
                if result_dadata_leasee()[0]['data']['opf']['short'] in ['ИП', 'КФХ', 'ГКФХ']:
                    deystvuysh_list_leasee = 'действующая'
            except:
                deystvuysh_list_leasee = 'действующего'
            imenyemoe = 'именуемое'
        return put_padezh_podpisant_rg, deystvuysh_list_leasee, imenyemoe, doverka_ustav_leasee

    def replace():
        eq_val = equipment_valute()
        # logging.info(f'{eq_val=}')
        data_xlsx = read_xlsx(path_application, pl)  # все из xlsx
        price_entry = data_xlsx[15]  # цена ПЛ
        # logging.info(f'{price_entry=}')
        suma_dann = number_to_words(price_entry, currency)

        payment_dkp = payment_for_dkp(price_entry)  # все для порядка оплаты
        info_about_seller = result_dadata()
        # logging.info(f'{info_about_seller=}')
        info_about_seller_director = some_info_seller(info_about_seller)
        # logging.info(f'{info_about_seller_director=}')
        full_seller = full_rekviti_seller(info_about_seller)
        dbase = DADATA_BASE.clean("name", full_seller[-2])
        rod_padezh_seller = dbase['result_genitive']
        name_and_dover_seller = seller_dkp_all()
        what_gender = dbase['gender']
        deistv_sell = deistv_seller(info_about_seller, full_seller[-2], what_gender)
        kratk_name_seller = result_dadata()[0]['data']['name']['short_with_opf']
        ident_lkmb_rt = identification_lkmb_rt(signatory, investor)
        ident_pl = indentification_pl(currency)
        id_ls = identification_leasee(data_xlsx[18])
        req1 = data_xlsx[10]
        ogrn1 = data_xlsx[13]
        add_inf = addicted_info_leasee(req1, ogrn1)

        months = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня',
                  7: 'июля', 8: 'августа', 9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}

        old_words_dkp = ["{{ new_old_pl }}", "{{ pb_vizor }}", "{{ identif_punkt_3_1_1 }}", "{{ identif_punkt_3_1_3 }}",
                         "{{ punkt_3_1_9 }}", "{{ punkt_3_3_3_key }}", "{{ punkt_3_3_3_key2 }}", "{{ punkt_3_1_6 }}",
                         "{{ punkt_3_3_7 }}", "{{ punkt_8_2 }}", "{{ pril_1_2 }}", "{{ pril_1_3 }}",
                         "{{ punkt_4_7_1 }}", "{{ punkt_5_3 }}", "{{ punkt_6_4 }}", "{{ punkt_8_3 }}",
                         "{{ type_currency }}", "{{ equivalent_currency }}", "{{ punkt_2_3_1_pay }}",
                         "{{ punkt_2_3_2_pay }}", "{{ punkt_2_3_3_pay }}", "{{ punkt_2_3_4_pay }}",
                         "{{ punkt_2_3_5_pay }}", "{{ punkt_2_3_1_num }}", "{{ punkt_2_3_2_num }}",
                         "{{ punkt_2_3_3_num }}", "{{ punkt_2_3_4_num }}", "{{ punkt_2_3_5_num }}",
                         "{{ punkt_2_3_1_numb_pr }}", "{{ punkt_2_3_2_numb_pr }}", "{{ punkt_2_3_3_numb_pr }}",
                         "{{ punkt_2_3_4_numb_pr }}", "{{ punkt_2_3_5_numb_pr }}", "{{ payment_1 }}", "{{ payment_2 }}",
                         "{{ payment_3 }}", "{{ payment_4 }}", "{{ payment_5 }}", "{{ payment_1_propis }}",
                         "{{ payment_2_propis }}", "{{ payment_3_propis }}", "{{ payment_4_propis }}",
                         "{{ payment_5_propis }}", "{{ full_name_seller }}", "{{ imenyemoe_dkp }}",
                         "{{ leader_seller_rod_padezh }}",
                         "{{ put_padezh_podpisant_seller }}",
                         "{{ deystvuysh_list_seller }}", "{{ doverka_ustav_seller }}", "{{ FULL_KRAKT_NAME_SELLER }}",
                         "{{ leader_seller_rod }}",
                         "{{ formatted_name_leader_seller }}", "{{ address_seller_dkp }}",
                         "{{ inn_kpp_seller }}", "{{ ogrn_seller }}",
                         "{{ rekvizit_leasee_bik }}", "{{ rekvizit_leasee_cs_shet }}", "{{ rekvizit_leasee_shet }}",
                         "{{ rekvizit_leasee_bank }}", "{{ main_activity_leasee }}", "{{ fio_leader }}",
                         "{{ email_leasee }}", "{{ phone_leasee }}", "{{ FULL_KRAKT_NAME_LEASEE }}",
                         "{{ ustav_capital }}", "{{ date_regist }}", "{{ okpo_leasee }}",
                         "{{ okato_leasee }}", "{{ ogrn_leasee }}",
                         # ниже список скорее всего не нужен, надо попробовать убрать его
                         "{{ inn_seller_list }}",
                         # 1 цена предмета лизинга, #2 предмет лизинга
                         "{{ price_entry }}", "{{ pl_entry }}",
                         "{{ formatted_name_leader_leasee }}",
                         "{{ leader_leasee }}", "{{ address_leasee_expluatazia }}", "{{ address_leasee }}",
                         "{{ inn_kpp_leasee }}", "{{ full_name_leasee }}",
                         # ниже будет подписант и прочие данные ЛКМБ-РТ
                         "{{ a_lkmb }}", "{{ lkmb_podpisant }}", "{{ preambula_dolj_lkmb }}",
                         "{{ preambula_fio_lkmb }}", "{{ deystvuysh }}", "{{ doverka_ustav }}",
                         "{{ r_chet_lkmb }}", "{{ bank_rekv_lkmb }}", "{{ kor_chet_lkmb }}", "{{ bik_lkmb }}",
                         "{{ currency_test }}",
                         # даты в ДКП
                         "{{ dt.today().day }}", "{{ months[dt.today().month] }}", "{{ dt.today().year }}",
                         "{{ number_dl }}",
                         "{{ suma_dann[0] }}",
                         # данные по лизингополучателю
                         "{{ leader_leasee_pod }}", "{{ leader_leasee_rod_padezh }}", "{{ imenyemoe }}",
                         "{{ put_padezh_podpisant_rg }}", "{{ deystvuysh_list_leasee }}",
                         "{{ doverka_ustav_leasee }}",
                         ]

        try:
            payment_dkp[15] = f'{float(payment_dkp[15]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[15] = ''
        try:
            payment_dkp[16] = f'{float(payment_dkp[16]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[16] = ''
        try:
            payment_dkp[17] = f'{float(payment_dkp[17]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[17] = ''
        try:
            payment_dkp[18] = f'{float(payment_dkp[18]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[18] = ''
        try:
            payment_dkp[19] = f'{float(payment_dkp[19]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[19] = ''

        new_words_dkp = [str(eq_val[0]), str(eq_val[1]), str(eq_val[2]), str(eq_val[3]), str(eq_val[4]),
                         str(eq_val[5]), str(eq_val[6]), str(eq_val[7]), str(eq_val[8]), str(eq_val[9]),
                         str(eq_val[10]), str(eq_val[11]), str(eq_val[12]), str(eq_val[13]), str(eq_val[14]),
                         str(eq_val[15]), str(eq_val[16]), str(eq_val[17]),
                         str(payment_dkp[0]), str(payment_dkp[1]), str(payment_dkp[2]), str(payment_dkp[3]),
                         str(payment_dkp[4]), str(payment_dkp[5]), str(payment_dkp[6]), str(payment_dkp[7]),
                         str(payment_dkp[8]), str(payment_dkp[9]), str(payment_dkp[10]), str(payment_dkp[11]),
                         str(payment_dkp[12]), str(payment_dkp[13]), str(payment_dkp[14]), payment_dkp[15],
                         payment_dkp[16], payment_dkp[17], payment_dkp[18], payment_dkp[19],
                         str(payment_dkp[20]), str(payment_dkp[21]), str(payment_dkp[22]), str(payment_dkp[23]),
                         str(payment_dkp[24]), str(name_and_dover_seller[0]), str(deistv_sell[1]),
                         str(info_about_seller_director[1]),
                         str(rod_padezh_seller), str(deistv_sell[0]),
                         str(name_and_dover_seller[1]), kratk_name_seller,
                         str(info_about_seller_director[0]), str(full_seller[4]), str(full_seller[6]),
                         str(full_seller[1]), str(full_seller[2]),
                         data_xlsx[0], data_xlsx[1], data_xlsx[2],
                         data_xlsx[3], data_xlsx[4], data_xlsx[5], data_xlsx[6], data_xlsx[7], data_xlsx[8],
                         data_xlsx[9], data_xlsx[10], data_xlsx[11], data_xlsx[12], data_xlsx[13],
                         # ниже ИНН поставщиков, цена и предмет лизинга
                         inn_seller,
                         f'{data_xlsx[15]:,.2f}'.replace(',', ' ').replace('.', ','),
                         pl,
                         data_xlsx[17], data_xlsx[18], data_xlsx[19], data_xlsx[20],
                         data_xlsx[21], data_xlsx[22],
                         # ниже будет подписант и прочие данные ЛКМБ-РТ
                         ident_lkmb_rt[0], ident_lkmb_rt[1], ident_lkmb_rt[2], ident_lkmb_rt[3], ident_lkmb_rt[5],
                         ident_lkmb_rt[4], ident_lkmb_rt[6], ident_lkmb_rt[7], ident_lkmb_rt[8], ident_lkmb_rt[9],
                         ident_pl[0],
                         # даты в ДКП
                         str(dt.today().day), str(months[dt.today().month]), str(dt.today().year),
                         str(numb_dl_dkp), str(suma_dann),
                         # данные по лизингополучателю
                         str(id_ls[1]), str(id_ls[0]),
                         str(add_inf[2]), str(add_inf[0]), str(add_inf[1]), str(add_inf[3])
                         ]

        return old_words_dkp, new_words_dkp

    def replace_words_in_dkp(docx_file, old_words_dkp, new_words_dkp):
        eq_val = equipment_valute()
        data_xlsx = read_xlsx(path_application, pl)
        # logging.info(f'{data_xlsx=}')
        price_entry = data_xlsx[15]
        payment_dkp = payment_for_dkp(price_entry)  # все для порядка оплаты
        info_about_seller = result_dadata()
        kratk_name_seller = result_dadata()[0]['data']['name']['short_with_opf']
        full_seller = full_rekviti_seller(info_about_seller)
        doc = Document(docx_file)
        print(f'Длина списка со старыми словами: {len(old_words_dkp)}')
        print(f'Длина списка с новыми словами: {len(new_words_dkp)}')
        for paragraph in doc.paragraphs:
            for i in range(len(old_words_dkp)):
                if old_words_dkp[i] in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_words_dkp[i], str(new_words_dkp[i]))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for i in range(len(old_words_dkp)):
                        if old_words_dkp[i] in cell.text:
                            cell.text = cell.text.replace(old_words_dkp[i], str(new_words_dkp[i]))

        if pnr == 'Нет':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '2.2. Сумма Договора включает стоимость доставки предмета лизинга до места эксплуатации и проведения пуско-наладочных работ.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '2.2. Сумма Договора не включает стоимость доставки предмета лизинга до места эксплуатации.']:
                    doc.element.body.remove(run._element)

        if equipment_or_not == 'None':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '- Техническую и товаросопроводительную документацию, в т.ч. технические паспорта, инструкции по запуску предмета лизинга в эксплуатацию и его последующему обслуживанию, чертежи общих видов с указанием габаритных размеров и характеристик; электрические, пневматические схемы; список запчастей и иную техническую и эксплуатационную документацию на русском языке, в том числе содержащую требования для его ввода в эксплуатацию;',
                    '- Сертификаты и/или иные документы, необходимые для постановки предмета лизинга на государственный учет (если применимо), и связанные с последующей его эксплуатацией;',
                    '- Счет(а) на оплату и счета-фактуры, выписанных в российских рублях.;',
                    '- Товарную накладную, выписанную в российских рублях и составленную по форме Госкомстата Российской Федерации ТОРГ-12.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in ['- счет(а) на оплату и счета-фактуры, выписанных в российских рублях;',
                                        '- Товарная накладная, выписанная в российских рублях и составленная по форме Госкомстата Российской Федерации ТОРГ-12;',
                                        '- копия грузовой таможенной декларации (в отношении автотранспортных средств, ввезенных на таможенную территорию России);',
                                        '- копия Декларации о товарах;',
                                        '- паспорт транспортного средства/самоходной машины/ выписка из ЭПТС;',
                                        '- сервисная книжка;',
                                        '- руководство по эксплуатации. ']:
                    doc.element.body.remove(run._element)

        if pnr == 'Нет':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '3.1.6. Произвести пуско-наладочные работы в соответствии с условиями настоящего договора.',
                    '3.1.7. Контролировать исполнение Лизингополучателем требований к производственному помещению согласно Техническому заданию (Приложение №1 к настоящему договору).',
                    '3.1.8. Письменно уведомить Покупателя о прикреплении Лизингополучателя к сервисному центру и своевременном прохождении технического обслуживания предмета лизинга.',
                    '3.2.3. Принять пуско-наладочные работы по акту приема пуско-наладочных работ. ',
                    '3.3.2. Принять пуско-наладочные работы по акту приема пуско-наладочных работ;',
                    '3.3.3. Обеспечить технические требования для проведения пуско-наладочных работ при условии соблюдения Продавцом п. 3.1.4 настоящего договора;',
                    '3.3.4. Соблюдать требования по эксплуатации предмета лизинга, установленные технической документацией переданной Продавцом, в соответствии с п. 3.1.3. настоящего Договора:',
                    '•	Использовать предмет лизинга только по назначению;',
                    '•	Использовать только соответствующий инструмент (расходный материал) на предмет лизинга;',
                    '•	Соблюдать иные условия, предусмотренные в технической документации на предмет лизинга. ']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '3.3.2. Соблюдать требования по эксплуатации предмета лизинга, установленные технической документацией переданной Продавцом, в соответствии с п. 3.1.3 настоящего Договора:',
                    '•	Использовать Предмет лизинга только по назначению;',
                    '•	Соблюдать иные условия, предусмотренные в технической документации на Предмет лизинга.',
                    f'3.3.3. В случае передачи Продавцом оригиналов документов {eq_val[5]}Лизингополучателю, в нарушении пункта 3.1.3 настоящего договора, Лизингополучатель обязан передать все оригиналы {eq_val[6]}Покупателю не позднее 1 (одного) рабочего дня нарочно, либо направить посредством срочной (курьерской) почты.',
                    '3.3.4. Прибыть к месту поставки предмета лизинга в установленный Продавцом срок в соответствии с п. 3.1.1 настоящего Договора.']:
                    doc.element.body.remove(run._element)

        if equipment_or_not == 'None':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    'Внести информацию о смене собственника предмета лизинга в Системе Электронных Паспортов (СЭП) не позднее дня подписания акта приема-передачи предмета лизинга, путём создания заявления на смену собственника и подтверждения его в СЭП квалифицированной электронной подписью Продавца. Фактом извещения Покупателя об исполненной обязанности является электронное сообщение с сайта СЭП на адрес электронной почты Покупателя указанный при регистрации в СЭП в виде сообщения с кодом для подтверждения Заявления на смену собственника. ']:
                    doc.element.body.remove(run._element)

        if place == "продавец" and acts == 'эксплуатация':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'4.1. Продавец передает Лизингополучателю, на основании выданной Покупателем доверенности,  предмет лизинга со своего склада по адресу: {full_seller[-1]}. Продавец указывает местонахождение склада в уведомлении о готовности предмета лизинга к передаче. Местом поставки является склад Продавца. Транспортировка предмета лизинга до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора, осуществляется за счет и силами Лизингополучателя со склада Продавца. Лизингополучатель обязуется застраховать предмет лизинга с момента отгрузки на время перевозки до момента поставки к месту эксплуатации.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания Продавцом, Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту поставки.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю со дня передачи предмета лизинга по товарной накладной и по акту приема-передачи в соответствии с условиями настоящего договора, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по месту отгрузки предмета лизинга со склада Продавца и подписания товарной накладной (ТОРГ-12) либо УПД.',
                    '4.1. Поставка предмета лизинга осуществляется за счет и силами Продавца до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания  Продавцом, Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту эксплуатации.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по акту приема-передачи в соответствии с условиями настоящего договора.']:
                    doc.element.body.remove(run._element)
        elif place == "продавец" and acts != 'эксплуатация':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'4.1. Транспортировка предмета лизинга до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора, осуществляется за счет и силами Лизингополучателя со склада Продавца по адресу: {full_seller[-1]}. Лизингополучатель обязуется застраховать предмет лизинга на время перевозки с момента отгрузки до момента поставки к месту эксплуатации. Местом поставки является место эксплуатации предмета лизинга.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания Продавцом,  Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту эксплуатации.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю со дня передачи предмета лизинга по акту  приема-передачи в соответствии с условиями настоящего договора, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по месту отгрузки предмета лизинга со склада Продавца и подписания товарной накладной (ТОРГ-12) либо УПД.',
                    '4.1. Поставка предмета лизинга осуществляется за счет и силами Продавца до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания  Продавцом, Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту эксплуатации.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по акту приема-передачи в соответствии с условиями настоящего договора.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'4.1. Продавец передает Лизингополучателю, на основании выданной Покупателем доверенности,  предмет лизинга со своего склада по адресу: {full_seller[-1]}. Продавец указывает местонахождение склада в уведомлении о готовности предмета лизинга к передаче. Местом поставки является склад Продавца. Транспортировка предмета лизинга до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора, осуществляется за счет и силами Лизингополучателя со склада Продавца. Лизингополучатель обязуется застраховать предмет лизинга с момента отгрузки на время перевозки до момента поставки к месту эксплуатации.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания Продавцом, Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту поставки.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю со дня передачи предмета лизинга по товарной накладной и по акту приема-передачи в соответствии с условиями настоящего договора, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по месту отгрузки предмета лизинга со склада Продавца и подписания товарной накладной (ТОРГ-12) либо УПД.',
                    f'4.1. Транспортировка предмета лизинга до места эксплуатации предмета лизинга по адресу, указанному в п.1.7. настоящего договора, осуществляется за счет и силами Лизингополучателя со склада Продавца по адресу: {full_seller[-1]}. Лизингополучатель обязуется застраховать предмет лизинга на время перевозки с момента отгрузки до момента поставки к месту эксплуатации. Местом поставки является место эксплуатации предмета лизинга.',
                    '4.2. Датой поставки предмета лизинга считается дата подписания Продавцом,  Покупателем и Лизингополучателем трехстороннего акта приема-передачи предмета лизинга по месту эксплуатации.',
                    '4.3. Право собственности на поставляемый предмет лизинга переходит от Продавца к Покупателю со дня передачи предмета лизинга по акту  приема-передачи в соответствии с условиями настоящего договора, а риск случайной гибели или случайного повреждения предмета лизинга, утраты, ущерба переходит от Продавца к Лизингополучателю со дня передачи предмета лизинга по месту отгрузки предмета лизинга со склада Продавца и подписания товарной накладной (ТОРГ-12) либо УПД.']:
                    doc.element.body.remove(run._element)

        if equipment_or_not == 'None':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '4.8. Лизингополучатель обязан осуществить предварительную подготовку помещения согласно Техническому заданию (Приложение №1 к настоящему договору), в котором будут проводиться монтажные и пуско-наладочные работы, обеспечить подключение водоснабжения и электроснабжения в течение 30 (тридцати) календарных дней с момента подписания настоящего договора.',
                    'В случае неисполнения обязательства, указанного в настоящем пункте, Покупатель вправе потребовать уплаты Лизингополучателем штрафа в размере 0,1 % (Ноль целых одна десятая процента) от стоимости предмета лизинга за каждый календарный день просрочки.',
                    '4.9.	Пуско-наладочные работы производятся силами Продавца с участием представителя Лизингополучателя по месту эксплуатации предмета лизинга. При проведении пуско-наладочных работ имеет право присутствовать представитель Покупателя. По результатам работ Продавец совместно с Лизингополучателем и Покупателем проводят проверку соответствия предмета лизинга техническим параметрам. Результаты проверки отражаются в акте приемки пуско-наладочных работ, подписываемом Продавцом, Покупателем и Лизингополучателем. Пуско-наладочные работы, включая приемку пуско-наладочных работ по акту, производятся в срок не позднее 8 календарных дней с момента поставки предмета лизинга. В случае несоответствия технических параметров Лизингополучатель вправе предъявить требования согласно п. 5.4. настоящего договора. ']:
                    doc.element.body.remove(run._element)

        if pl_new_or_not == 'new':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    'Поставляемый предмет лизинга принадлежит ему на праве собственности, никому другому не продан, под запретом (арестом), в залоге не состоит, не является предметом спора и свободен от каких-либо притязаний третьих лиц (не обременен правами третьих лиц (аренда, безвозмездное пользование и т.д.)), и право собственности на него никем не оспаривается. Предмет лизинга не находится в залоге у продавца до его полной оплаты.',
                    '5.2. Продавец передает предмет лизинга в комплектации соответствующей договору в порядке, установленном в п.4.1. настоящего договора.',
                    '5.3. Гарантийный срок на предмет лизинга устанавливается заводом производителем. Гарантийное обслуживание осуществляется на специализированных СТО, указанных в сервисной книжке к автомобилю.',
                    '5.1.1. Поставляемый предмет лизинга соответствует требованиям действующих в Российской Федерации стандартов, нормативно-технической документации, техническим характеристикам и параметрам, предусмотренным  настоящим договором;',
                    '5.1.2. Поставляемый предмет лизинга принадлежит ему на праве собственности, никому другому не продан, под запретом (арестом), в залоге не состоит, не является предметом спора и свободен от каких-либо притязаний третьих лиц (не обременен правами третьих лиц (аренда, безвозмездное пользование и т.д.)), и право собственности на него никем не оспаривается. Предмет лизинга не находится в залоге у продавца до его полной оплаты.',
                    '5.1.3. Прилагаемые инструкции по  эксплуатации являются комплектными и достаточными для эксплуатации предмета лизинга.',
                    '5.3. Лизингополучатель уведомлен, что гарантийный срок на предмет лизинга истек, Продавец не несет ответственности за неисправность предмета лизинга.',
                    '5.4. В случае выявления существенных дефектов предмета лизинга (в том числе, несоответствие техническим характеристикам, и т.д.), по письменному заявлению Лизингополучателя Покупатель вправе потребовать расторжения договора в судебном порядке.']:
                    doc.element.body.remove(run._element)
        elif pl_new_or_not == 'used_with_garantee':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '5.1.1. Поставляемый предмет лизинга соответствует требованиям действующих в Российской Федерации стандартов, нормативно-технической документации, техническим характеристикам и параметрам, предусмотренным настоящим договором;',
                    '5.1.2. Производительность поставляемого предмета лизинга и качество его работы полностью соответствуют техническим условиям, указанным в сопроводительной документации, то есть предмет лизинга, будет без дефектов, будет обеспечивать стабильное качество выполняемых с его помощью работ (соответствующей параметрам, указанным заводом-изготовителем в сопроводительной документации), и при нормальной эксплуатации не будет выявлено дефектов применяемых материалов и качества исполнения;',
                    '5.1.3. Прилагаемые инструкции по эксплуатации являются комплектными и достаточными для эксплуатации предмета лизинга.',
                    f'5.3. Гарантийное обслуживание предмета лизинга осуществляется Продавцом на сертифицированных заводом-изготовителем технических сервисах и/или аккредитованным изготовителем персоналом. Гарантийный срок на предмет лизинга устанавливается в соответствии с условиями завода-изготовителя, установленным паспортом или иной технической документацией, и составляет срок 36 (тридцать шесть) календарных месяцев с момента подписания акта приема-передачи предмета лизинга {eq_val[-6]}. Гарантийный срок предмета лизинга продлевается на время проведения ремонтных работ.',
                    '5.4. В случае выявления дефектов в период гарантийного срока эксплуатации предмета лизинга Лизингополучатель обязан в течение 5 (Пяти) рабочих дней с момента их обнаружения известить Продавца об обнаруженных дефектах посредством почты, факсимильной связи / электронной почты (при этом в претензии указываются реквизиты настоящего договора, наименования Продавца и Лизингополучателя, наименование предмета лизинга, выявленные дефекты, исполнитель и его контактный телефон). Выезд специалиста Продавца для проведения гарантийного ремонта производится за счет Продавца в течение 3 (Трех) рабочих дней с момента поступления заявки от Лизингополучателя. Продавец осуществляет замену предмета лизинга или его части в случае его наличия на складе в течение 10 (Десяти) рабочих дней с момента поступления претензии. В случае отсутствия необходимых частей на складе, Продавец обязуется в течение 30 (Тридцати) дней произвести замену предмета лизинга или его части.',
                    '5.5. В случае выявления существенных дефектов предмета лизинга (в том числе, несоответствие техническим характеристикам, и т.д.) по письменному заявлению Лизингополучателя Покупатель вправе потребовать расторжения договора.',
                    '5.6.  После окончания срока действия гарантии на предмет лизинга, ремонт предмета лизинга и поставка запасных частей осуществляется Продавцом за счет Лизингополучателя на основании соответствующих договоров.',
                    '5.1.1. Поставляемый предмет лизинга соответствует требованиям действующих в Российской Федерации стандартов, нормативно-технической документации, техническим характеристикам и параметрам, предусмотренным  настоящим договором;',
                    '5.1.2. Поставляемый предмет лизинга принадлежит ему на праве собственности, никому другому не продан, под запретом (арестом), в залоге не состоит, не является предметом спора и свободен от каких-либо притязаний третьих лиц (не обременен правами третьих лиц (аренда, безвозмездное пользование и т.д.)), и право собственности на него никем не оспаривается. Предмет лизинга не находится в залоге у продавца до его полной оплаты.',
                    '5.1.3. Прилагаемые инструкции по  эксплуатации являются комплектными и достаточными для эксплуатации предмета лизинга.',
                    '5.2. Продавец передает предмет лизинга в комплектации соответствующей договору в порядке, установленном в п.4.1. настоящего договора.',
                    '5.3. Лизингополучатель уведомлен, что гарантийный срок на предмет лизинга истек, Продавец не несет ответственности за неисправность предмета лизинга.',
                    '5.4. В случае выявления существенных дефектов предмета лизинга (в том числе, несоответствие техническим характеристикам, и т.д.), по письменному заявлению Лизингополучателя Покупатель вправе потребовать расторжения договора в судебном порядке.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '5.1.1. Поставляемый предмет лизинга соответствует требованиям действующих в Российской Федерации стандартов, нормативно-технической документации, техническим характеристикам и параметрам, предусмотренным настоящим договором;',
                    '5.1.2. Производительность поставляемого предмета лизинга и качество его работы полностью соответствуют техническим условиям, указанным в сопроводительной документации, то есть предмет лизинга, будет без дефектов, будет обеспечивать стабильное качество выполняемых с его помощью работ (соответствующей параметрам, указанным заводом-изготовителем в сопроводительной документации), и при нормальной эксплуатации не будет выявлено дефектов применяемых материалов и качества исполнения;',
                    '5.1.3. Прилагаемые инструкции по эксплуатации являются комплектными и достаточными для эксплуатации предмета лизинга.',
                    '5.2. Продавец передает предмет лизинга Покупателю и/или Лизингополучателю в комплектации соответствующей договору.',
                    f'5.3. Гарантийное обслуживание предмета лизинга осуществляется Продавцом на сертифицированных заводом-изготовителем технических сервисах и/или аккредитованным изготовителем персоналом. Гарантийный срок на предмет лизинга устанавливается в соответствии с условиями завода-изготовителя, установленным паспортом или иной технической документацией, и составляет срок 36 (тридцать шесть) календарных месяцев с момента подписания акта приема-передачи предмета лизинга {eq_val[-6]}. Гарантийный срок предмета лизинга продлевается на время проведения ремонтных работ.',
                    '5.4. В случае выявления дефектов в период гарантийного срока эксплуатации предмета лизинга Лизингополучатель обязан в течение 5 (Пяти) рабочих дней с момента их обнаружения известить Продавца об обнаруженных дефектах посредством почты, факсимильной связи / электронной почты (при этом в претензии указываются реквизиты настоящего договора, наименования Продавца и Лизингополучателя, наименование предмета лизинга, выявленные дефекты, исполнитель и его контактный телефон). Выезд специалиста Продавца для проведения гарантийного ремонта производится за счет Продавца в течение 3 (Трех) рабочих дней с момента поступления заявки от Лизингополучателя. Продавец осуществляет замену предмета лизинга или его части в случае его наличия на складе в течение 10 (Десяти) рабочих дней с момента поступления претензии. В случае отсутствия необходимых частей на складе, Продавец обязуется в течение 30 (Тридцати) дней произвести замену предмета лизинга или его части.',
                    '5.5. В случае выявления существенных дефектов предмета лизинга (в том числе, несоответствие техническим характеристикам, и т.д.) по письменному заявлению Лизингополучателя Покупатель вправе потребовать расторжения договора.',
                    '5.6.  После окончания срока действия гарантии на предмет лизинга, ремонт предмета лизинга и поставка запасных частей осуществляется Продавцом за счет Лизингополучателя на основании соответствующих договоров.',
                    'Поставляемый предмет лизинга принадлежит ему на праве собственности, никому другому не продан, под запретом (арестом), в залоге не состоит, не является предметом спора и свободен от каких-либо притязаний третьих лиц (не обременен правами третьих лиц (аренда, безвозмездное пользование и т.д.)), и право собственности на него никем не оспаривается. Предмет лизинга не находится в залоге у продавца до его полной оплаты.',
                    '5.3. Гарантийный срок на предмет лизинга устанавливается заводом производителем. Гарантийное обслуживание осуществляется на специализированных СТО, указанных в сервисной книжке к автомобилю.']:
                    doc.element.body.remove(run._element)

        if pl_new_or_not != 'new':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '6.6. В случае выявления неуплаты утилизационного сбора в отношении предмета лизинга в нарушение обязанности, возникшей по основаниям, имевшим место до передачи предмета лизинга Покупателю, Покупатель вправе потребовать от Продавца возмещения убытков, включая расходы, связанные с уплатой Покупателем сбора, сумм штрафов, а также любых иных расходов и (или) издержек, которые возникли у Покупателя в связи с тем, что сбор подлежал уплате до передачи предмета лизинга Покупателю, но не был уплачен.']:
                    doc.element.body.remove(run._element)

        if diadok == 'диадок':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '8.9.В рамках исполнения настоящего Договора, Стороны установили возможность использования Электронного документооборота (далее – ЭДО). Любой документ, представленный в электронно-цифровом формате, подписанный/заверенный действующей на момент передачи усиленной квалифицированной ЭЦП уполномоченного стороной лица и переданный в рамках Договора, является эквивалентом идентичного по содержанию документа на бумажном носителе, подписанного уполномоченным лицом организации-отправителя с проставлением печати, имеет равную с ним юридическую силу и порождает для Сторон аналогичные права и обязанности. При осуществлении обмена электронными документами Стороны используют форматы документов, утвержденные приказами ФНС России. Если форматы документов не утверждены, Стороны используют согласованные между собой форматы, при этом в ЭДО включаются следующие виды документов: УПД; УКД.',
                    'Сторона обязана подписать документ в электронно-цифровом формате в течение 3 (трёх) рабочих дней с даты выставления его другой Стороной договора. Датой выставления электронного документа по телекоммуникационным каналам связи считается дата подтверждения Оператором ЭДО выставления такого электронного документа. В случае неполучения Стороной подтверждения подписания выставленного документа, такой документ считается подписанным другой Стороной в указанный срок.',
                    'Стороны признают, что:',
                    '- несут ответственность за все электронные документы, оформленные и переданные в рамках настоящего договора,',
                    '- используемые средства подготовки, передачи и проверки электронных документов достаточны для обеспечения надежного, эффективного и безопасного документооборота,',
                    '- содержание электронного документа соответствует полномочиям лица, его подписавшего.',
                    'Каждая Сторона вправе приостановить ЭДО в случаях:',
                    '- обнаружения технических неисправностей своей автоматизированной системы ЭДО;',
                    '- несоблюдения одной из Сторон требований к ЭДО и обеспечению информационной безопасности, установленных законодательством РФ;',
                    '- изменения банковских и иных реквизитов, имеющих существенное значение для определения юридического статуса и идентификации Сторон через оператора ЭДО;',
                    '- по инициативе одной из Сторон, при письменном уведомлении не позднее 5 (пяти) рабочих дней до предполагаемой даты и срока приостановления.',
                    'На период приостановления ЭДО Стороны переходят на бумажный документооборот, порядок и сроки которого согласованы Сторонами в рамках Договора. Возобновление ЭДО производится на основании письменного уведомления Стороной – инициатором приостановления ЭДО другой Стороны не позднее 5 (пяти) рабочих дней до предполагаемой даты возобновления ЭДО. ЭДО возобновляется в назначенный срок Стороной не позднее 1 (одного) рабочего дня до назначенной даты возобновления ЭДО.',
                    'По всем вопросам, не оговоренным настоящим разделом Договора, Стороны руководствуются Гражданским кодексом РФ, Налоговым кодексом РФ, Федеральным законом от 06.04.2011 №63-ФЗ «Об электронной подписи», Федеральным законом от 06.11.2011 №402—ФЗ «О бухгалтерском учете» и иными действующими нормативно-правовыми актами РФ.']:
                    doc.element.body.remove(run._element)
        elif diadok == 'эцп':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '8.9. Стороны соглашаются осуществлять документооборот при исполнении настоящего Договора в электронном виде по телекоммуникационным каналам связи, в том числе через систему оператора электронного документооборота «Диадок» (далее - Диадок) с использованием квалифицированной электронной подписи уполномоченного стороной лица. На каждом электронном документе (сообщении), отправляемом через Диадок, автоматически проставляется отметка о передаче документа с указанием даты и времени. Электронные документы, подписанные квалифицированной электронной подписью, признаются электронными документами, равнозначными документам на бумажном носителе, подписанными собственноручной подписью и заверенным печатью, за исключением случая, если федеральными законами или принимаемыми в соответствии с ними нормативными правовыми актами установлено требование о необходимости составления документа исключительно на бумажном носителе. Сторона обязана подписать документ в электронно-цифровом формате не позднее 2(двух) рабочих дней с даты выставления его другой Стороной договора. В случае неполучения Стороной подтверждения подписания выставленного документа, такой документ считается подписанным другой Стороной в указанный срок.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    '8.9.В рамках исполнения настоящего Договора, Стороны установили возможность использования Электронного документооборота (далее – ЭДО). Любой документ, представленный в электронно-цифровом формате, подписанный/заверенный действующей на момент передачи усиленной квалифицированной ЭЦП уполномоченного стороной лица и переданный в рамках Договора, является эквивалентом идентичного по содержанию документа на бумажном носителе, подписанного уполномоченным лицом организации-отправителя с проставлением печати, имеет равную с ним юридическую силу и порождает для Сторон аналогичные права и обязанности. При осуществлении обмена электронными документами Стороны используют форматы документов, утвержденные приказами ФНС России. Если форматы документов не утверждены, Стороны используют согласованные между собой форматы, при этом в ЭДО включаются следующие виды документов: УПД; УКД.',
                    'Сторона обязана подписать документ в электронно-цифровом формате в течение 3 (трёх) рабочих дней с даты выставления его другой Стороной договора. Датой выставления электронного документа по телекоммуникационным каналам связи считается дата подтверждения Оператором ЭДО выставления такого электронного документа. В случае неполучения Стороной подтверждения подписания выставленного документа, такой документ считается подписанным другой Стороной в указанный срок.',
                    'Стороны признают, что:',
                    '- несут ответственность за все электронные документы, оформленные и переданные в рамках настоящего договора,',
                    '- используемые средства подготовки, передачи и проверки электронных документов достаточны для обеспечения надежного, эффективного и безопасного документооборота,',
                    '- содержание электронного документа соответствует полномочиям лица, его подписавшего.',
                    'Каждая Сторона вправе приостановить ЭДО в случаях:',
                    '- обнаружения технических неисправностей своей автоматизированной системы ЭДО;',
                    '- несоблюдения одной из Сторон требований к ЭДО и обеспечению информационной безопасности, установленных законодательством РФ;',
                    '- изменения банковских и иных реквизитов, имеющих существенное значение для определения юридического статуса и идентификации Сторон через оператора ЭДО;',
                    '- по инициативе одной из Сторон, при письменном уведомлении не позднее 5 (пяти) рабочих дней до предполагаемой даты и срока приостановления.',
                    'На период приостановления ЭДО Стороны переходят на бумажный документооборот, порядок и сроки которого согласованы Сторонами в рамках Договора. Возобновление ЭДО производится на основании письменного уведомления Стороной – инициатором приостановления ЭДО другой Стороны не позднее 5 (пяти) рабочих дней до предполагаемой даты возобновления ЭДО. ЭДО возобновляется в назначенный срок Стороной не позднее 1 (одного) рабочего дня до назначенной даты возобновления ЭДО.',
                    'По всем вопросам, не оговоренным настоящим разделом Договора, Стороны руководствуются Гражданским кодексом РФ, Налоговым кодексом РФ, Федеральным законом от 06.04.2011 №63-ФЗ «Об электронной подписи», Федеральным законом от 06.11.2011 №402—ФЗ «О бухгалтерском учете» и иными действующими нормативно-правовыми актами РФ.',
                    '8.9. Стороны соглашаются осуществлять документооборот при исполнении настоящего Договора в электронном виде по телекоммуникационным каналам связи, в том числе через систему оператора электронного документооборота «Диадок» (далее - Диадок) с использованием квалифицированной электронной подписи уполномоченного стороной лица. На каждом электронном документе (сообщении), отправляемом через Диадок, автоматически проставляется отметка о передаче документа с указанием даты и времени. Электронные документы, подписанные квалифицированной электронной подписью, признаются электронными документами, равнозначными документам на бумажном носителе, подписанными собственноручной подписью и заверенным печатью, за исключением случая, если федеральными законами или принимаемыми в соответствии с ними нормативными правовыми актами установлено требование о необходимости составления документа исключительно на бумажном носителе. Сторона обязана подписать документ в электронно-цифровом формате не позднее 2(двух) рабочих дней с даты выставления его другой Стороной договора. В случае неполучения Стороной подтверждения подписания выставленного документа, такой документ считается подписанным другой Стороной в указанный срок.']:
                    doc.element.body.remove(run._element)

        if eq_val[-1] == 'Обычная оплата':
            for run in doc.paragraphs:

                if run.text.strip() in [
                    '2.3. Оплата суммы договора производится в рублях по курсу обмена валют ЦБ РФ на день платежа на расчетный счет Продавца и осуществляется в следующем порядке:',
                    '2.3. Оплата суммы договора производится в рублях по средневзвешенному курсу до 4-х знаков после запятой CNY/RUB (Китайский юань / Российский рубль) ПАО «Московская Биржа» (https://www.moex.com/ru/issue/CNYRUB_TOM/CETS) предшествующего дня на расчетный счет Продавца и осуществляется в следующем порядке:',
                    '2.4. Сторонами, на следующий рабочий день, после дня оплаты, производится пересчет полученной суммы оплаты в рублях по средневзвешенному курсу CNY/RUB (Китайский юань / Российский рубль) ПАО «Московская Биржа» (https://www.moex.com/ru/issue/CNYRUB_TOM/CETS), действующему на день оплаты.',
                    'По возникшей разнице в юанях будет производиться перерасчет между Покупателем и Продавцом, в следующем порядке:',
                    '- если, средневзвешенный курс, определенный на день оплаты выше, чем курс, по которому Покупатель производил оплату, то Покупатель обязан в течение 1 рабочего дня доплатить Продавцу разницу;',
                    '- если средневзвешенный курс, определенный на день оплаты ниже, чем курс, по которому Покупатель производил оплату, то Продавец обязан в течение 1 рабочего дня вернуть Покупателю разницу.',
                    'Пересчет суммы оплаты в рублях производится только единоразово в рабочий день, следующий за днем оплаты.',
                    'Требование   настоящего  пункта,  связанное  с  возможным переносом  срока  оплаты,  не  являются  основанием  для  признания оплаты Предмета лизинга Покупателем с просрочкой и не может служить основанием для привлечения Покупателя  к  ответственности.']:
                    doc.element.body.remove(run._element)
        elif eq_val[-1] == 'moex':
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3. Оплата суммы договора производится в {eq_val[16]} и осуществляется в следующем порядке:',
                    '2.3. Оплата суммы договора производится в рублях по курсу обмена валют ЦБ РФ на день платежа на расчетный счет Продавца и осуществляется в следующем порядке:',
                    '2.4. Платеж считается произведенным в момент списания денежных средств с расчетного счета Покупателя.']:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3. Оплата суммы договора производится в {eq_val[16]} и осуществляется в следующем порядке:',
                    '2.3. Оплата суммы договора производится в рублях по средневзвешенному курсу до 4-х знаков после запятой CNY/RUB (Китайский юань / Российский рубль) ПАО «Московская Биржа» (https://www.moex.com/ru/issue/CNYRUB_TOM/CETS) предшествующего дня на расчетный счет Продавца и осуществляется в следующем порядке:',
                    '2.4. Сторонами, на следующий рабочий день, после дня оплаты, производится пересчет полученной суммы оплаты в рублях по средневзвешенному курсу CNY/RUB (Китайский юань / Российский рубль) ПАО «Московская Биржа» (https://www.moex.com/ru/issue/CNYRUB_TOM/CETS), действующему на день оплаты.',
                    'По возникшей разнице в юанях будет производиться перерасчет между Покупателем и Продавцом, в следующем порядке:',
                    '- если, средневзвешенный курс, определенный на день оплаты выше, чем курс, по которому Покупатель производил оплату, то Покупатель обязан в течение 1 рабочего дня доплатить Продавцу разницу;',
                    '- если средневзвешенный курс, определенный на день оплаты ниже, чем курс, по которому Покупатель производил оплату, то Продавец обязан в течение 1 рабочего дня вернуть Покупателю разницу.',
                    'Пересчет суммы оплаты в рублях производится только единоразово в рабочий день, следующий за днем оплаты.',
                    'Требование   настоящего  пункта,  связанное  с  возможным переносом  срока  оплаты,  не  являются  основанием  для  признания оплаты Предмета лизинга Покупателем с просрочкой и не может служить основанием для привлечения Покупателя  к  ответственности.'
                ]:
                    doc.element.body.remove(run._element)

        try:
            payment_dkp[15] = f'{float(payment_dkp[15]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[15] = ''
        try:
            payment_dkp[16] = f'{float(payment_dkp[16]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[16] = ''
        try:
            payment_dkp[17] = f'{float(payment_dkp[17]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[17] = ''
        try:
            payment_dkp[18] = f'{float(payment_dkp[18]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[18] = ''
        try:
            payment_dkp[19] = f'{float(payment_dkp[19]):,.2f}'.replace(',', ' ').replace('.', ',')
        except ValueError:
            payment_dkp[19] = ''

        if len(payment_order.split(' ')) == 1:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3.1. {payment_dkp[0]} платеж в размере {payment_dkp[5]}% ({payment_dkp[10]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[15]} {payment_dkp[20]}, в т.ч. НДС 20%, в течение 15 (пятнадцати) рабочих дней с момента перечисления Лизингополучателем первоначального платежа Покупателю по договору лизинга, а также после выставления Продавцом счета на оплату.',
                    f'2.3.2. {payment_dkp[1]} платеж в размере {payment_dkp[6]}% ({payment_dkp[11]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[16]} {payment_dkp[21]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней с момента получения уведомления о готовности предмета лизинга к отгрузке согласно п.3.1.1 и после выставления Продавцом счета на оплату.',
                    f'2.3.3. {payment_dkp[2]} платеж в размере {payment_dkp[7]}% ({payment_dkp[12]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[17]} {payment_dkp[22]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.',
                    f'2.3.4. {payment_dkp[3]} платеж в размере {payment_dkp[8]}% ({payment_dkp[13]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[18]} {payment_dkp[23]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.',
                    f'2.3.5. {payment_dkp[4]} платеж в размере {payment_dkp[9]}% ({payment_dkp[14]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[19]} {payment_dkp[24]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.']:
                    doc.element.body.remove(run._element)

        elif len(payment_order.split(' ')) == 2:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3.1. Оплата в размере {payment_dkp[5]}% ({payment_dkp[10]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[15]} {payment_dkp[20]}, в т.ч. НДС, в течение 15 (пятнадцати) рабочих дней с момента перечисления Лизингополучателем первоначального платежа Покупателю по договору лизинга,',
                    'и в течение 7 (семи) рабочих дней с момента получения Покупателем уведомления Продавца о готовности предмета лизинга к отгрузке согласно п.3.1.1 и после выставления Продавцом счета на оплату.',
                    f'2.3.3. {payment_dkp[2]} платеж в размере {payment_dkp[7]}% ({payment_dkp[12]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[17]} {payment_dkp[22]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.',
                    f'2.3.4. {payment_dkp[3]} платеж в размере {payment_dkp[8]}% ({payment_dkp[13]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[18]} {payment_dkp[23]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.',
                    f'2.3.5. {payment_dkp[4]} платеж в размере {payment_dkp[9]}% ({payment_dkp[14]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[19]} {payment_dkp[24]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.'
                ]:
                    doc.element.body.remove(run._element)
        elif len(payment_order.split(' ')) == 3:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3.1. Оплата в размере {payment_dkp[5]}% ({payment_dkp[10]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[15]} {payment_dkp[20]}, в т.ч. НДС, в течение 15 (пятнадцати) рабочих дней с момента перечисления Лизингополучателем первоначального платежа Покупателю по договору лизинга,',
                    'и в течение 7 (семи) рабочих дней с момента получения Покупателем уведомления Продавца о готовности предмета лизинга к отгрузке согласно п.3.1.1 и после выставления Продавцом счета на оплату.',
                    f'2.3.4. {payment_dkp[3]} платеж в размере {payment_dkp[8]}% ({payment_dkp[13]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[18]} {payment_dkp[23]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.',
                    f'2.3.5. {payment_dkp[4]} платеж в размере {payment_dkp[9]}% ({payment_dkp[14]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[19]} {payment_dkp[24]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.'
                ]:
                    doc.element.body.remove(run._element)
        elif len(payment_order.split(' ')) == 4:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3.1. Оплата в размере {payment_dkp[5]}% ({payment_dkp[10]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[15]} {payment_dkp[20]}, в т.ч. НДС, в течение 15 (пятнадцати) рабочих дней с момента перечисления Лизингополучателем первоначального платежа Покупателю по договору лизинга,',
                    'и в течение 7 (семи) рабочих дней с момента получения Покупателем уведомления Продавца о готовности предмета лизинга к отгрузке согласно п.3.1.1 и после выставления Продавцом счета на оплату.',
                    f'2.3.5. {payment_dkp[4]} платеж в размере {payment_dkp[9]}% ({payment_dkp[14]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[19]} {payment_dkp[24]}, в т.ч. НДС 20%, в течение 7 (семи) рабочих дней после выставления Продавцом счета на оплату.'
                ]:
                    doc.element.body.remove(run._element)
        else:
            for run in doc.paragraphs:
                if run.text.strip() in [
                    f'2.3.1. Оплата в размере {payment_dkp[5]}% ({payment_dkp[10]}) от стоимости предмета лизинга, что составляет сумму {eq_val[-2]} {payment_dkp[15]} {payment_dkp[20]}, в т.ч. НДС, в течение 15 (пятнадцати) рабочих дней с момента перечисления Лизингополучателем первоначального платежа Покупателю по договору лизинга,',
                    'и в течение 7 (семи) рабочих дней с момента получения Покупателем уведомления Продавца о готовности предмета лизинга к отгрузке согласно п.3.1.1 и после выставления Продавцом счета на оплату.']:
                    doc.element.body.remove(run._element)

        full_krakt_name_leasee = data_xlsx[8].replace('"', '')
        dir_path = Path('webapp') / 'static' / 'agreements' / f'{full_krakt_name_leasee} {inn_client}' / f'{dt.today().strftime(f"%d.%m.%Y")}'
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
        doc.save(dir_path / fr"ДКП {inn_client}.docx")

    print(equipment_or_not)
    if equipment_or_not is None:
        replace_words_in_dkp(dkp_car_ooo_path, replace()[0],
                             replace()[1])
    else:
        replace_words_in_dkp(dkp_obor_ooo_path, replace()[0],
                             replace()[1])
